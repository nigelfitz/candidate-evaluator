# Candidate Analyser
from __future__ import annotations

import io, os, re, json, hashlib, textwrap, sys, html
from dataclasses import dataclass
from typing import List, Dict, Tuple, Optional, Any

import numpy as np
import pandas as pd
import streamlit as st

# -----------------------------
# Page Configuration (must be first)
# -----------------------------
st.set_page_config(
    page_title="Candidate Analyser",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS to reduce top padding and add app branding
st.markdown("""
    <style>
    /* Reduce top padding */
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    /* App branding header with distinct styling */
    header[data-testid="stHeader"] {
        background-color: rgba(255, 255, 255, 0.95);
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #f0f2f6;
    }
    header[data-testid="stHeader"]::before {
        content: "🎯 Candidate Analyser";
        font-size: 1.5rem;
        font-weight: 700;
        padding-left: 1rem;
        color: #1f77b4;
        white-space: nowrap;
        display: inline-block;
        letter-spacing: 0.5px;
    }
    /* Add spacing after header before main content */
    .main .block-container {
        margin-top: 1rem;
    }
    /* Center text vertically in info boxes */
    .stAlert {
        padding-top: 0.9rem !important;
        padding-bottom: 0.9rem !important;
    }
    </style>
""", unsafe_allow_html=True)

# -----------------------------
# API Key detection (global)
# -----------------------------
API_KEY_SET = bool(os.getenv('OPENAI_API_KEY'))

def extract_text_layout_aware(file_bytes: bytes, filename: str) -> str:
    """
    Extract PDF text in a layout-aware order using PyMuPDF (fitz):
    - Read page text as blocks with coordinates
    - Sort by y (top) then x (left) to respect columns/boxes
    - Join blocks with blank lines
    Falls back to empty string on error; caller should handle fallback.
    """
    try:
        if fitz is None:
            return ""
        text_blocks = []
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                blocks = page.get_text("blocks") or []
                blocks.sort(key=lambda b: (round(b[1]), round(b[0])))
                for b in blocks:
                    t = (b[4] or "").strip()
                    if t:
                        text_blocks.append(t)
        return "\n\n".join(text_blocks).strip()
    except Exception:
        return ""


# -----------------------------
# Optional / lazy imports
# -----------------------------
def _try_import(name: str):
    try:
        return __import__(name)
    except Exception:
        return None

pdfplumber = _try_import("pdfplumber")
docx = _try_import("docx")
sklearn = _try_import("sklearn")
dotenv = _try_import("dotenv")
reportlab = _try_import("reportlab")
_tabulate = _try_import("tabulate")
openpyxl = _try_import("openpyxl")

# OpenAI SDK (support both modern client and legacy)
_openai_mod = _try_import("openai")
try:
    # New SDK style
    from openai import OpenAI  # type: ignore
except Exception:
    OpenAI = None  # type: ignore

# Robust PDF extractor (PyMuPDF) and optional OCR fallback
fitz = _try_import("fitz")  # PyMuPDF
pdf2image = _try_import("pdf2image")
pytesseract = _try_import("pytesseract")
PIL_mod = _try_import("PIL")
Image = getattr(PIL_mod, "Image", None) if PIL_mod else None

if sklearn:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity as sk_cosine
else:
    TfidfVectorizer = None
    sk_cosine = None

st_models = _try_import("sentence_transformers")
if st_models:
    from sentence_transformers import SentenceTransformer
    import torch
else:
    SentenceTransformer = None
    torch = None

if reportlab:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.units import cm

if dotenv is not None:
    try:
        from dotenv import load_dotenv
        load_dotenv(override=False)
    except Exception:
        pass

# -----------------------------
# Data classes
# -----------------------------
@dataclass
class Candidate:
    name: str
    file_name: str
    text: str
    hash: str
    raw_bytes: Optional[bytes] = None  # keep original bytes for PDF viewing

@dataclass
class JD:
    file_name: str
    text: str
    hash: str
    raw_bytes: Optional[bytes] = None  # keep original bytes for robust extractor if needed

@dataclass
class JDSections:
    key_skills: List[str]
    responsibilities: List[str]
    qualifications: List[str]
    experience_required: List[str]

def _sections_to_dict(secs: JDSections) -> Dict[str, list]:
    return {
        "key_skills": list(secs.key_skills or []),
        "responsibilities": list(secs.responsibilities or []),
        "qualifications": list(secs.qualifications or []),
        "experience_required": list(secs.experience_required or []),
    }
# -----------------------------
# Hash helpers
# -----------------------------
def hash_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()

def hash_text(s: str) -> str:
    return hashlib.sha256(s.encode("utf-8", errors="ignore")).hexdigest()

# -----------------------------
# OpenAI helpers
# -----------------------------
def _get_openai_client():
    """
    Returns (mode, client_or_module):
      - ("modern", OpenAI()) if openai>=1.0 client is available
      - ("legacy", openai) if old ChatCompletion API is available
      - (None, None) if nothing usable
    """
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    
    if not api_key:
        return (None, None)
    # Prefer modern client
    if OpenAI is not None:
        try:
            client = OpenAI(api_key=api_key)
            return ("modern", client)
        except Exception:
            pass
    # Try legacy module
    if _openai_mod is not None:
        try:
            _openai_mod.api_key = api_key
            return ("legacy", _openai_mod)
        except Exception:
            pass
    return (None, None)

# -----------------------------
# Cost Estimation & Usage Logging
# -----------------------------
def estimate_tokens(text: str) -> int:
    """Rough token estimation: ~4 chars per token for English text."""
    return len(text) // 4

def estimate_analysis_cost(jd_text: str, candidate_texts: List[str], num_criteria: int, model: str = "gpt-4o") -> Dict[str, Any]:
    """
    Estimate the cost of analyzing candidates.
    
    Returns:
        dict with keys: total_input_tokens, total_output_tokens, estimated_cost_usd, breakdown
    """
    # Pricing per 1M tokens (as of Nov 2024)
    PRICING = {
        "gpt-4o": {"input": 2.50, "output": 10.00},
        "gpt-4o-mini": {"input": 0.15, "output": 0.60},
        "gpt-4-turbo": {"input": 10.00, "output": 30.00}
    }
    
    prices = PRICING.get(model, PRICING["gpt-4o"])
    
    # JD extraction (if using GPT)
    jd_tokens_in = estimate_tokens(jd_text)
    jd_tokens_out = num_criteria * 20  # ~20 tokens per criterion
    
    # Candidate insights
    total_candidate_tokens_in = sum(estimate_tokens(text) for text in candidate_texts)
    # Each candidate insight: JD + criteria + candidate text → detailed analysis
    avg_candidate_input = (jd_tokens_in + (num_criteria * 15) + (total_candidate_tokens_in // max(1, len(candidate_texts))))
    total_insight_input = avg_candidate_input * len(candidate_texts)
    total_insight_output = len(candidate_texts) * 500  # ~500 tokens per insight
    
    # Total tokens
    total_input = jd_tokens_in + total_insight_input
    total_output = jd_tokens_out + total_insight_output
    
    # Cost calculation
    cost_input = (total_input / 1_000_000) * prices["input"]
    cost_output = (total_output / 1_000_000) * prices["output"]
    total_cost = cost_input + cost_output
    
    return {
        "total_input_tokens": total_input,
        "total_output_tokens": total_output,
        "estimated_cost_usd": round(total_cost, 2),
        "cost_per_candidate": round(total_cost / max(1, len(candidate_texts)), 2),
        "breakdown": {
            "jd_extraction": round((jd_tokens_in / 1_000_000) * prices["input"] + (jd_tokens_out / 1_000_000) * prices["output"], 2),
            "candidate_insights": round(cost_input + cost_output - ((jd_tokens_in / 1_000_000) * prices["input"] + (jd_tokens_out / 1_000_000) * prices["output"]), 2)
        },
        "model": model
    }

def log_usage(
    num_candidates: int,
    num_criteria: int,
    total_resume_chars: int,
    jd_chars: int,
    estimated_cost: float,
    model: str = "gpt-4o",
    log_file: str = "usage_log.csv"
):
    """Log analysis usage to CSV file for cost tracking."""
    import csv
    from datetime import datetime
    
    log_entry = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "num_candidates": num_candidates,
        "num_criteria": num_criteria,
        "total_resume_chars": total_resume_chars,
        "jd_chars": jd_chars,
        "estimated_cost_usd": estimated_cost,
        "model": model
    }
    
    # Check if file exists to write header
    file_exists = os.path.exists(log_file)
    
    try:
        with open(log_file, "a", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=log_entry.keys())
            if not file_exists:
                writer.writeheader()
            writer.writerow(log_entry)
    except Exception as e:
        # Don't fail the analysis if logging fails
        st.warning(f"Could not log usage: {e}")

def call_llm_json(system_prompt: str, user_prompt: str, schema: Dict[str, Any], model: str = "gpt-4o") -> Dict[str, Any]:
    """
    Calls OpenAI to return structured JSON adhering to 'schema'.
    Works with both the modern and legacy python SDK variants.
    """
    mode, cli = _get_openai_client()
    if mode is None or cli is None:
        # No API available
        return {"key_skills": [], "responsibilities": [], "qualifications": [], "experience_required": []}

    # Build a compact instruction to strongly bias JSON-only output
    json_guard = (
        "Return ONLY valid JSON that conforms to the provided JSON Schema. "
        "Do not include backticks or any commentary."
    )

    if mode == "modern":
        # function-free JSON via response_format in new SDK
        try:
            resp = cli.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"{user_prompt}\n\n{json_guard}\n\nSchema:\n{json.dumps(schema)}"}
                ],
                temperature=0.2,
                response_format={"type": "json_object"},
            )
            txt = resp.choices[0].message.content or "{}"
            return json.loads(txt)
        except Exception:
            return {"key_skills": [], "responsibilities": [], "qualifications": [], "experience_required": []}
    else:
        # legacy: best-effort JSON
        try:
            resp = cli.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"{user_prompt}\n\n{json_guard}\n\nSchema:\n{json.dumps(schema)}"}
                ],
                temperature=0.2,
            )
            txt = resp["choices"][0]["message"]["content"] or "{}"
            return json.loads(txt)
        except Exception:
            return {"key_skills": [], "responsibilities": [], "qualifications": [], "experience_required": []}

# -----------------------------
# File reading / extraction
# -----------------------------
@st.cache_data(show_spinner=False)
def read_file_bytes(file_bytes: bytes, file_name: str) -> str:
    lower = file_name.lower()
    if lower.endswith(".txt"):
        try:
            return file_bytes.decode("utf-8")
        except Exception:
            return file_bytes.decode("latin-1", errors="ignore")
    if lower.endswith(".docx") and docx is not None:
        try:
            d = docx.Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in d.paragraphs)
        except Exception:
            pass
    if lower.endswith(".pdf") and pdfplumber is not None:
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
            return "\n".join(pages)
        except Exception:
            pass
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text_from_pdf_robust(file_bytes: bytes, strong: bool = False, prefer_pdf2image: bool = True) -> str:
    """
    Primary: PyMuPDF text extraction.
    Fallback: rasterize and OCR with pytesseract if very low yield.
    Uses pdf2image if available; otherwise rasterizes via PyMuPDF to avoid Poppler dependency.
    If `strong=True`, render at higher DPI and use stronger thresholding to recover faint text.
    """
    text = ""
    # ---- Pass 1: direct text extraction via PyMuPDF ----
    if fitz is not None:
        try:
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                chunks = []
                for p in doc:
                    chunks.append(p.get_text("text"))
                text = "\n".join(chunks)
        except Exception:
            text = ""

    # ---- Pass 2: OCR when text is too small ----
    if len((text or "").strip()) < 200 and pytesseract is not None:
        ocr_chunks = []
        try:
            dpi = 400 if strong else 300
            thresh_cut = 170 if strong else 200  # lower cut -> darker letters kept
            tess_cfg = "--oem 3 --psm 6"
            if pdf2image is not None:
                from pdf2image import convert_from_bytes
                pages = convert_from_bytes(file_bytes, dpi=dpi)
                for img in pages:
                    if Image is not None:
                        gray = img.convert("L")
                        bw = gray.point(lambda x: 0 if x < thresh_cut else 255, "1").convert("L")
                        img_for_ocr = bw
                    else:
                        img_for_ocr = img
                    ocr_chunks.append(pytesseract.image_to_string(img_for_ocr, config=tess_cfg))
            elif fitz is not None:
                with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                    mat = fitz.Matrix(dpi/72.0, dpi/72.0)
                    for page in doc:
                        pix = page.get_pixmap(matrix=mat, alpha=False)
                        if Image is not None:
                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                            gray = img.convert("L")
                            bw = gray.point(lambda x: 0 if x < thresh_cut else 255, "1").convert("L")
                            ocr_chunks.append(pytesseract.image_to_string(bw, config=tess_cfg))
                        else:
                            ocr_chunks.append("")
            text = "\n".join([t for t in ocr_chunks if t])
        except Exception:
            pass

    return (text or "").strip()# -----------------------------
# Text utils
# -----------------------------
def normalize_ws(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()

# -----------------------------
# Embeddings / scoring
# -----------------------------
@st.cache_resource(show_spinner=False)
def load_embedder() -> Dict[str, Any]:
    if SentenceTransformer is not None:
        try:
            model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
            device = "cuda" if (torch and torch.cuda.is_available()) else "cpu"
            model.to(device)
            return {"type": "sbert", "model": model, "device": device}
        except Exception:
            pass
    return {"type": "tfidf"}

def embed_sbert(texts: List[str], model) -> np.ndarray:
    return np.array(model.encode(texts, normalize_embeddings=True, show_progress_bar=False))

def pairwise_cosine(a: np.ndarray, b: np.ndarray) -> np.ndarray:
    a_norm = a / (np.linalg.norm(a, axis=1, keepdims=True) + 1e-12)
    b_norm = b / (np.linalg.norm(b, axis=1, keepdims=True) + 1e-12)
    return a_norm @ b_norm.T

def chunk_text(text: str, max_chars: int = 1200, overlap: int = 150) -> List[str]:
    text = normalize_ws(text)
    if len(text) <= max_chars:
        return [text]
    chunks, start = [], 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = max(0, end - overlap)
    return chunks

# -----------------------------
# GPT JD extraction (new)
# -----------------------------
def _normalize_lines(items: List[str]) -> List[str]:
    out = []
    for it in items or []:
        t = re.sub(r"\s+", " ", (it or "").strip(" •-–·\t")).strip()
        if t:
            out.append(t)
    # de-dupe while preserving order
    seen = set()
    deduped = []
    for x in out:
        xl = x.lower()
        if xl not in seen:
            seen.add(xl)
            deduped.append(x)
    return deduped

def extract_jd_sections_with_gpt(jd_text: str) -> JDSections:
    schema = {
        "type": "object",
        "properties": {
            "key_skills": {"type": "array", "items": {"type": "string"}},
            "responsibilities": {"type": "array", "items": {"type": "string"}},
            "qualifications": {"type": "array", "items": {"type": "string"}},
            "experience_required": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["key_skills", "responsibilities", "qualifications", "experience_required"],
        "additionalProperties": False,
    }
    system = (
        "You extract criteria from job descriptions into clear bullet items. "
        "Be concise, de-duplicate similar bullets, and keep each item atomic."
    )
    user = (
        "Extract the following JD into four sections (key_skills, responsibilities, "
        "qualifications, experience_required). Return ONLY JSON matching the provided schema.\n\n"
        f"JD:\n{jd_text}"
    )
    data = call_llm_json(system, user, schema)
    return JDSections(
        key_skills=_normalize_lines(data.get("key_skills", [])),
        responsibilities=_normalize_lines(data.get("responsibilities", [])),
        qualifications=_normalize_lines(data.get("qualifications", [])),
        experience_required=_normalize_lines(data.get("experience_required", [])),
    )

@st.cache_resource()

def cached_extract_jd_sections_with_gpt(jd_hash: str, jd_text: str) -> JDSections:
    # Cache key includes jd_hash so repeated runs on the same JD don't re-call GPT
    if 'last_cached_jd_hash' in st.session_state and st.session_state['last_cached_jd_hash'] == jd_hash:
        try: st.toast('JD extracted from cache ✅', icon='💾')
        except Exception: pass
    else:
        st.session_state['last_cached_jd_hash'] = jd_hash
    return extract_jd_sections_with_gpt(jd_text)

def build_criteria_from_gpt_sections(sections: JDSections, per_section: int = 6, cap_total: int = 30) -> Tuple[List[str], Dict[str,str]]:
    order = [
        ("key_skills", sections.key_skills),
        ("responsibilities", sections.responsibilities),
        ("qualifications", sections.qualifications),
        ("experience_required", sections.experience_required),
    ]
    crits: List[str] = []
    cat_map: Dict[str, str] = {}
    for sec_name, items in order:
        picked = items[:per_section]
        for p in picked:
            if p and p not in crits:
                crits.append(p); cat_map[p] = sec_name
            if len(crits) >= cap_total:
                break
        if len(crits) >= cap_total:
            break
    return crits[:cap_total], cat_map

# -----------------------------
# Embedding & analysis (unchanged)
# -----------------------------
@st.cache_data(show_spinner=False)
def prepare_corpus_embeddings(corpus_chunks: List[str], _embedder_info: Dict[str, Any]):
    info = _embedder_info
    if info.get("type") == "sbert" and info.get("model") is not None:
        embs = embed_sbert(corpus_chunks, info["model"])
        return embs, None
    if TfidfVectorizer is not None and sk_cosine is not None:
        vec = TfidfVectorizer(max_features=5000, ngram_range=(1, 2))
        mat = vec.fit_transform(corpus_chunks)
        return mat, vec
    vocab: Dict[str, int] = {}
    rows: List[Dict[str, int]] = []
    for doc in corpus_chunks:
        counts: Dict[str, int] = {}
        for t in re.findall(r"\b\w+\b", (doc or "").lower()):
            counts[t] = counts.get(t, 0) + 1
        rows.append(counts)
        for t in counts:
            if t not in vocab:
                vocab[t] = len(vocab)
    mat = np.zeros((len(rows), len(vocab)), dtype=float)
    for i, row in enumerate(rows):
        for t, c in row.items():
            mat[i, vocab[t]] = c
    return mat, {"type": "bow", "vocab": vocab}

def infer_candidate_name(file_name: str, text: str) -> str:
    lines = [l.strip() for l in (text or "").splitlines()][:60]
    for l in lines:
        m = re.match(r"(?i)^\s*name\s*[:\-]\s*(.+)$", l)
        if m:
            cand = normalize_ws(m.group(1))
            if 2 <= len(cand) <= 60 and not re.search(r"\d|@|http", cand):
                return cand.title()
    def looks_like_name(s: str) -> bool:
        if not s or len(s) > 60:
            return False
        if re.search(r"\d|@|http", s):
            return False
        parts = [p for p in re.split(r"\s+", s) if p]
        if len(parts) < 2 or len(parts) > 4:
            return False
        bad = {"curriculum", "vitae", "resume", "cv", "profile", "professional", "summary"}
        if any(w.lower() in bad for w in parts):
            return False
        cap_like = sum(1 for p in parts if re.match(r"^[A-Z][a-z'\-]+$", p))
        return cap_like >= max(2, len(parts) - 1)
    for l in lines:
        if looks_like_name(l):
            return l.title()
    base = os.path.splitext(os.path.basename(file_name))[0]
    base = re.sub(r"(?i)\b(cv|resume|curriculum|vitae)\b", " ", base)
    base = re.sub(r"[_\-\.]+", " ", base)
    base = re.sub(r"\d+", " ", base)
    base = normalize_ws(base)
    parts = base.split()
    if 2 <= len(parts) <= 4:
        return " ".join(p.capitalize() for p in parts)
    return (parts[0].capitalize() if parts else "Candidate")

def compute_max_similarity_to_chunks(query_texts, chunk_embs, info, vec_or_meta):
    if info["type"] == "sbert":
        q_embs = embed_sbert(query_texts, info["model"])
        sims = pairwise_cosine(q_embs, chunk_embs)
        argmax = sims.argmax(axis=1)
        return sims.max(axis=1), argmax
    if vec_or_meta is None:
        return np.zeros(len(query_texts)), np.zeros(len(query_texts), dtype=int)
    if hasattr(vec_or_meta, "transform"):
        q_mat = vec_or_meta.transform(query_texts)
        sims = sk_cosine(q_mat, chunk_embs)
        argmax = np.asarray(sims.argmax(axis=1)).ravel()
        return sims.max(axis=1), argmax
    vocab = vec_or_meta["vocab"]
    def bow_vec(s: str):
        v = np.zeros((len(vocab),), dtype=float)
        for t in re.findall(r"\b\w+\b", (s or "").lower()):
            if t in vocab:
                v[vocab[t]] += 1
        return v
    q_mat = np.stack([bow_vec(s) for s in query_texts], axis=0)
    denom = (np.linalg.norm(q_mat, axis=1, keepdims=True) + 1e-12) * (np.linalg.norm(chunk_embs, axis=1, keepdims=True).T + 1e-12)
    sims = (q_mat @ chunk_embs.T) / denom
    argmax = sims.argmax(axis=1)
    return sims.max(axis=1), argmax

def analyse_candidates(
    candidates: List[Candidate],
    criteria: List[str],
    weights: Optional[List[float]] = None,
    chunk_chars: int = 1200,
    overlap: int = 150,
    _progress_parent=None
) -> Tuple[pd.DataFrame, Dict[str, Dict], Dict[str, List[Tuple[str, float]]], Dict[Tuple[str,str], Tuple[str,float]]]:
    criteria = [normalize_ws(c) for c in (criteria or []) if normalize_ws(c)]
    if not criteria or not candidates:
        return pd.DataFrame(), {}, {}, {}

    if weights is None or len(weights) != len(criteria):
        weights = [1.0] * len(criteria)
    weights = np.array(weights, dtype=float)
    weights = weights / (weights.sum() if weights.sum() > 0 else len(weights))

    all_chunk_texts, chunk_index = [], []
    for i, c in enumerate(candidates):
        chunks = chunk_text(c.text, max_chars=chunk_chars, overlap=overlap)
        for j, ch in enumerate(chunks):
            all_chunk_texts.append(ch); chunk_index.append((i, j))

    info = load_embedder()
    chunk_embs, vec_or_meta = prepare_corpus_embeddings(all_chunk_texts, info)

    cand_to_rows = {i: [] for i in range(len(candidates))}
    for ridx, (ci, cj) in enumerate(chunk_index):
        cand_to_rows[ci].append(ridx)

    coverage_records, insights_map, snippets_map = [], {}, {}
    evidence_map: Dict[Tuple[str,str], Tuple[str,float]] = {}

    progress = (_progress_parent.progress(0.0, text="Scoring…") if _progress_parent is not None else None)
    if progress is None:
        class _NoProgress:
            def progress(self, *args, **kwargs):
                return self
            def empty(self):
                pass
        progress = _NoProgress()
    total = max(1, len(candidates))

    for i, cand in enumerate(candidates):
        rows = cand_to_rows[i]
        cand_chunks = [all_chunk_texts[r] for r in rows]
        if len(cand_chunks) == 0:
            crit_scores = np.zeros(len(criteria)); argmax = np.zeros(len(criteria), dtype=int)
        else:
            cand_chunk_embs = chunk_embs[rows] if not isinstance(chunk_embs, np.ndarray) else chunk_embs[rows, :]
            crit_scores, argmax = compute_max_similarity_to_chunks(criteria, cand_chunk_embs, info, vec_or_meta)

        overall = float((crit_scores * weights).sum())
        row = {"Candidate": cand.name, "Overall": overall}
        for c_name, s in zip(criteria, crit_scores):
            row[c_name] = float(s)
        coverage_records.append(row)

        order = np.argsort(-crit_scores)
        top = [f"{criteria[idx]} — strong alignment ({crit_scores[idx]:.2f})" for idx in order[:3]]
        gaps = [f"{criteria[idx]} — weak coverage ({crit_scores[idx]:.2f})" for idx in order[-3:]]
        notes = f"Overall weighted score: {overall:.2f}. Doc length: {len(cand.text):,} chars."
        insights_map[cand.name] = {"top": top, "gaps": gaps, "notes": notes}

        top_snips = []
        for qi, best_idx in enumerate(argmax):
            if cand_chunks:
                ch = cand_chunks[int(best_idx)]
                top_snips.append((ch[:600] + ("…" if len(ch) > 600 else ""), float(crit_scores[qi])))
                evidence_map[(cand.name, criteria[qi])] = (ch[:800], float(crit_scores[qi]))
        seen, uniq_snips = set(), []
        for snip, sc in sorted(top_snips, key=lambda x: -x[1]):
            key = snip[:80]
            if key not in seen:
                uniq_snips.append((snip, sc)); seen.add(key)
            if len(uniq_snips) >= 5: break
        snippets_map[cand.name] = uniq_snips

        progress.progress((i + 1) / total, text=f"Scoring {cand.name}…")

    try:
        progress.progress(1.0, text="Scoring complete.")
        progress.empty()
    except Exception:
        pass

    coverage_df = pd.DataFrame(coverage_records)
    if not coverage_df.empty:
        coverage_df.sort_values(by="Overall", ascending=False, inplace=True)
        coverage_df.reset_index(drop=True, inplace=True)

    return coverage_df, insights_map, snippets_map, evidence_map

# -----------------------------
# Criteria parse & weights
# -----------------------------
def parse_criteria_text(criteria_text: str) -> List[str]:
    items = [normalize_ws(x) for x in (criteria_text or "").splitlines()]
    return [x for x in items if x]

def get_weights(criteria: List[str], mode: str, weights_csv: str) -> List[float]:
    if mode == "Uniform" or not criteria:
        return [1.0] * len(criteria)
    lines = [l.strip() for l in (weights_csv or "").splitlines() if l.strip()]
    if not lines:
        return [1.0] * len(criteria)
    mapping, vals = {}, []
    for line in lines:
        if "," in line:
            k, v = line.split(",", 1)
            try: mapping[normalize_ws(k)] = float(v)
            except Exception: pass
        else:
            try: vals.append(float(line))
            except Exception: pass
    if mapping:
        return [float(mapping.get(c, 1.0)) for c in criteria]
    if vals and len(vals) == len(criteria):
        return [float(x) for x in vals]
    return [1.0] * len(criteria)

# -----------------------------
# Report export
# -----------------------------
def _md_table_or_fallback(df: pd.DataFrame) -> str:
    if _tabulate is None:
        cols = list(df.columns)
        lines = ["| " + " | ".join(cols) + " |", "| " + " | ".join(["---"] * len(cols)) + " |"]
        for _, row in df.iterrows():
            lines.append("| " + " | ".join(str(row[c]) for c in cols) + " |")
        return "\n".join(lines)
    return df.to_markdown(index=False)

def to_markdown_report(coverage: pd.DataFrame, insights: Dict[str, Dict], jd_text: str) -> str:
    if coverage.empty:
        return "# Candidate Report\n\n_No analysis available._"
    lines = ["# Candidate Analysis Report", ""]
    lines.append("## Job Description (excerpt)")
    jd_excerpt = (jd_text or "").strip().splitlines()
    jd_excerpt = "\n".join(jd_excerpt[:20])
    lines.append("" if not jd_excerpt else jd_excerpt)
    lines.append("")
    lines.append("## Leaderboard (Overall)")
    top_tbl = coverage[["Candidate", "Overall"]].copy()
    top_tbl["Overall"] = top_tbl["Overall"].map(lambda x: f"{x:.2f}")
    lines.append(_md_table_or_fallback(top_tbl))
    lines.append("")
    lines.append("## Coverage Matrix")
    cov_tbl = coverage.copy()
    for c in cov_tbl.columns:
        if c != "Candidate":
            cov_tbl[c] = cov_tbl[c].map(lambda x: f"{x:.2f}")
    lines.append(_md_table_or_fallback(cov_tbl))
    lines.append("")
    lines.append("## Candidate Insights")
    for name, info in insights.items():
        lines.append(f"### {name}")
        lines.append("**Top strengths**:")
        for t in info.get("top", []): lines.append(f"- {t}")
        lines.append("**Gaps / risks**:")
        for g in info.get("gaps", []): lines.append(f"- {g}")
        notes = info.get("notes", "")
        if notes: lines.append(f"_Notes:_ {notes}")
        lines.append("")
    return "\n".join(lines)

def to_pdf_bytes_from_markdown(md_text: str) -> Optional[bytes]:
    if reportlab is None: return None
    buf = io.BytesIO(); c = rl_canvas.Canvas(buf, pagesize=A4)
    left = 2.0 * cm; top = A4[1] - 2.0 * cm
    wrapped = []
    for para in (md_text or "").splitlines():
        wrapped.extend(textwrap.wrap(para, width=90) or [""])
    y = top
    for line in wrapped:
        if y < 2.0 * cm:
            c.showPage(); y = top
        c.drawString(left, y, line); y -= 12
    c.save(); buf.seek(0)
    return buf.getvalue()

# -------- Phase 1 Export Functions --------

def to_excel_coverage_matrix(coverage: pd.DataFrame, cat_map: Dict[str, str], hi: float = 0.70, lo: float = 0.45) -> Optional[bytes]:
    """Export coverage matrix to Excel with formatting and color coding."""
    if openpyxl is None:
        return None
    
    try:
        from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
        from openpyxl.utils.dataframe import dataframe_to_rows
        
        buf = io.BytesIO()
        
        # Create workbook with multiple sheets
        wb = openpyxl.Workbook()
        
        # Sheet 1: Summary
        ws_summary = wb.active
        ws_summary.title = "Summary"
        
        # Add title
        ws_summary['A1'] = "Candidate Analysis Summary"
        ws_summary['A1'].font = Font(size=16, bold=True, color="1F77B4")
        ws_summary.merge_cells('A1:D1')
        
        # Add metrics
        row = 3
        ws_summary[f'A{row}'] = "Total Candidates:"
        ws_summary[f'B{row}'] = len(coverage)
        ws_summary[f'A{row}'].font = Font(bold=True)
        
        row += 1
        crit_cols = [c for c in coverage.columns if c not in ('Candidate', 'Overall')]
        ws_summary[f'A{row}'] = "Total Criteria:"
        ws_summary[f'B{row}'] = len(crit_cols)
        ws_summary[f'A{row}'].font = Font(bold=True)
        
        row += 2
        ws_summary[f'A{row}'] = "Top 5 Candidates"
        ws_summary[f'A{row}'].font = Font(size=14, bold=True)
        row += 1
        
        # Top candidates table
        top5 = coverage[['Candidate', 'Overall']].head(5)
        ws_summary[f'A{row}'] = "Rank"
        ws_summary[f'B{row}'] = "Candidate"
        ws_summary[f'C{row}'] = "Overall Score"
        for col in ['A', 'B', 'C']:
            ws_summary[f'{col}{row}'].font = Font(bold=True)
            ws_summary[f'{col}{row}'].fill = PatternFill(start_color="E7F3FF", end_color="E7F3FF", fill_type="solid")
        
        for idx, (_, r) in enumerate(top5.iterrows(), 1):
            row += 1
            ws_summary[f'A{row}'] = idx
            ws_summary[f'B{row}'] = r['Candidate']
            ws_summary[f'C{row}'] = round(r['Overall'], 2)
            
            # Color code the score
            score = r['Overall']
            if score >= hi:
                ws_summary[f'C{row}'].fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
            elif score >= lo:
                ws_summary[f'C{row}'].fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
            else:
                ws_summary[f'C{row}'].fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
        
        # Adjust column widths
        ws_summary.column_dimensions['A'].width = 8
        ws_summary.column_dimensions['B'].width = 30
        ws_summary.column_dimensions['C'].width = 15
        
        # Sheet 2: Coverage Matrix
        ws_matrix = wb.create_sheet("Coverage Matrix")
        
        # Transpose the coverage matrix: Criteria as rows, Candidates as columns
        # Get criteria columns (exclude Candidate and Overall)
        criteria_cols = [c for c in coverage.columns if c not in ('Candidate', 'Overall')]
        
        # Create transposed data structure
        # First row: header with "Criterion" then candidate names
        header_row = ['Criterion'] + coverage['Candidate'].tolist()
        
        # Write header
        for c_idx, value in enumerate(header_row, 1):
            cell = ws_matrix.cell(row=1, column=c_idx, value=value)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="1F77B4", end_color="1F77B4", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")
        
        # Write data rows (one criterion per row)
        for r_idx, criterion in enumerate(criteria_cols, 2):
            # First column: criterion name
            cell = ws_matrix.cell(row=r_idx, column=1, value=criterion)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="E7F3FF", end_color="E7F3FF", fill_type="solid")
            cell.alignment = Alignment(horizontal="left")
            
            # Remaining columns: scores for each candidate
            for c_idx, (_, candidate_row) in enumerate(coverage.iterrows(), 2):
                score = candidate_row[criterion]
                cell = ws_matrix.cell(row=r_idx, column=c_idx, value=round(score, 2))
                
                # Color code the score
                if score >= hi:
                    cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
                elif score >= lo:
                    cell.fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
                else:
                    cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
                cell.alignment = Alignment(horizontal="center")
        
        # Auto-adjust column widths
        for column in ws_matrix.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws_matrix.column_dimensions[column_letter].width = adjusted_width
        
        # Sheet 3: Legend
        ws_legend = wb.create_sheet("Legend")
        ws_legend['A1'] = "Score Color Coding"
        ws_legend['A1'].font = Font(size=14, bold=True)
        
        ws_legend['A3'] = "Strong (≥0.70)"
        ws_legend['A3'].fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
        
        ws_legend['A4'] = "Moderate (0.45-0.70)"
        ws_legend['A4'].fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
        
        ws_legend['A5'] = "Weak (<0.45)"
        ws_legend['A5'].fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
        
        ws_legend.column_dimensions['A'].width = 25
        
        wb.save(buf)
        buf.seek(0)
        return buf.getvalue()
    except Exception as e:
        st.error(f"Excel export error: {e}")
        return None

def to_executive_summary_pdf(
    coverage: pd.DataFrame,
    insights: Dict[str, Dict],
    jd_text: str,
    cat_map: Dict[str, str],
    hi: float = 0.70,
    lo: float = 0.45,
    jd_filename: str = "Job Description"
) -> Optional[bytes]:
    """Generate a professional executive summary PDF."""
    if reportlab is None:
        return None
    
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from datetime import datetime
        
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1F77B4'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#1F77B4'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Title
        story.append(Paragraph("Executive Summary", title_style))
        story.append(Paragraph("Candidate Analysis Report", styles['Heading2']))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
        story.append(Spacer(1, 0.5*cm))
        
        # Job Description Summary
        story.append(Paragraph("Position", heading_style))
        
        # Extract job title from JD content (usually in first few lines)
        job_title = "Position Not Specified"
        jd_lines = [line.strip() for line in jd_text.split('\n') if line.strip()]
        
        # Look for common job title patterns in first 5 lines
        for line in jd_lines[:5]:
            # Skip very short lines, URLs, dates
            if len(line) < 10 or 'http' in line.lower() or any(char.isdigit() for char in line[:4]):
                continue
            # Look for title indicators
            if any(indicator in line.lower() for indicator in ['position:', 'role:', 'job title:', 'title:']):
                job_title = line.split(':', 1)[1].strip() if ':' in line else line
                break
            # Or just use first substantial line (likely the title)
            elif len(line) > 15 and len(line) < 100:
                job_title = line
                break
        
        story.append(Paragraph(f"<b>Position:</b> {job_title}", styles['Normal']))
        
        # Add key requirements summary (skip the title line, get next meaningful content)
        jd_summary_lines = []
        title_found = False
        for line in jd_lines:
            # Skip the title line
            if not title_found and job_title in line:
                title_found = True
                continue
            # Collect meaningful lines (not too short, not URLs/emails)
            if len(line) > 30 and 'http' not in line.lower() and '@' not in line:
                jd_summary_lines.append(line)
            if len(jd_summary_lines) >= 3:  # Get 3 good lines
                break
        
        if jd_summary_lines:
            summary_text = ' '.join(jd_summary_lines)[:400]  # Max 400 chars
            story.append(Paragraph(f"<b>Key Requirements:</b> {summary_text}...", styles['Normal']))
        
        story.append(Spacer(1, 0.3*cm))
        
        # Analysis Overview
        story.append(Paragraph("Analysis Overview", heading_style))
        overview_data = [
            ["Metric", "Value"],
            ["Total Candidates Analyzed", str(len(coverage))],
            ["Evaluation Criteria", str(len([c for c in coverage.columns if c not in ('Candidate', 'Overall')]))],
            ["Analysis Date", datetime.now().strftime('%Y-%m-%d')]
        ]
        overview_table = Table(overview_data, colWidths=[8*cm, 8*cm])
        overview_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F77B4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(overview_table)
        story.append(Spacer(1, 0.5*cm))
        
        # Top Candidates
        story.append(Paragraph("Top 5 Candidates", heading_style))
        top5 = coverage[['Candidate', 'Overall']].head(5)
        
        top_data = [["Rank", "Candidate Name", "Overall Score", "Rating"]]
        for idx, (_, row) in enumerate(top5.iterrows(), 1):
            score = row['Overall']
            rating = "Strong" if score >= hi else ("Moderate" if score >= lo else "Weak")
            top_data.append([str(idx), row['Candidate'], f"{score:.2f}", rating])
        
        top_table = Table(top_data, colWidths=[2*cm, 7*cm, 3*cm, 4*cm])
        top_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F77B4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
        ]))
        
        # Color code ratings
        for i in range(1, len(top_data)):
            score = top5.iloc[i-1]['Overall']
            if score >= hi:
                top_table.setStyle(TableStyle([('BACKGROUND', (3, i), (3, i), colors.HexColor('#C6EFCE'))]))
            elif score >= lo:
                top_table.setStyle(TableStyle([('BACKGROUND', (3, i), (3, i), colors.HexColor('#FFEB9C'))]))
            else:
                top_table.setStyle(TableStyle([('BACKGROUND', (3, i), (3, i), colors.HexColor('#FFC7CE'))]))
        
        story.append(top_table)
        story.append(Spacer(1, 0.5*cm))
        
        # Key Insights for Top 3
        story.append(Paragraph("Key Insights - Top 3 Candidates", heading_style))
        
        for idx, (_, row) in enumerate(top5.head(3).iterrows(), 1):
            name = row['Candidate']
            score = row['Overall']
            
            story.append(Paragraph(f"<b>{idx}. {name}</b> (Score: {score:.2f})", styles['Heading3']))
            
            if name in insights:
                info = insights[name]
                
                # Strengths
                strengths = info.get('top', [])
                if strengths:
                    story.append(Paragraph("<b>Strengths:</b>", styles['Normal']))
                    for s in strengths[:3]:  # Top 3 strengths
                        story.append(Paragraph(f"• {s}", styles['Normal']))
                
                # Gaps
                gaps = info.get('gaps', [])
                if gaps:
                    story.append(Paragraph("<b>Development Areas:</b>", styles['Normal']))
                    for g in gaps[:2]:  # Top 2 gaps
                        story.append(Paragraph(f"• {g}", styles['Normal']))
            else:
                story.append(Paragraph("<i>Detailed insights not available</i>", styles['Normal']))
            
            story.append(Spacer(1, 0.3*cm))
        
        # Recommendation (intelligent, context-aware)
        story.append(Paragraph("Recommendation", heading_style))
        
        top_candidate = coverage.iloc[0]
        top_score = top_candidate['Overall']
        
        # Check if we have multiple strong candidates
        strong_candidates = coverage[coverage['Overall'] >= hi]
        moderate_candidates = coverage[(coverage['Overall'] >= lo) & (coverage['Overall'] < hi)]
        weak_candidates = coverage[coverage['Overall'] < lo]
        
        # Build intelligent recommendation based on actual data
        recommendation_parts = []
        
        if len(strong_candidates) == 0:
            # No strong candidates at all
            if len(moderate_candidates) > 0:
                recommendation_parts.append(
                    f"<b>Caution:</b> No candidates achieved a strong match score (≥{hi:.2f}). "
                    f"The highest score was <b>{top_score:.2f}</b> for <b>{top_candidate['Candidate']}</b>, "
                    f"indicating a <b>moderate match</b>."
                )
                recommendation_parts.append(
                    f"We recommend carefully reviewing the {len(moderate_candidates)} moderate-scoring candidate(s) "
                    f"to identify specific skill gaps, or consider expanding the candidate pool."
                )
            else:
                recommendation_parts.append(
                    f"<b>Warning:</b> All candidates scored below the moderate threshold ({lo:.2f}). "
                    f"The highest score was only <b>{top_score:.2f}</b> for <b>{top_candidate['Candidate']}</b>."
                )
                recommendation_parts.append(
                    "We recommend reconsidering the job requirements or sourcing additional candidates, "
                    "as the current pool shows weak alignment with the position criteria."
                )
        
        elif len(strong_candidates) == 1:
            # Clear winner
            recommendation_parts.append(
                f"<b>{top_candidate['Candidate']}</b> is the clear leading candidate with a strong overall score "
                f"of <b>{top_score:.2f}</b>, demonstrating excellent alignment with the position requirements."
            )
            
            if len(moderate_candidates) > 0:
                recommendation_parts.append(
                    f"Additionally, {len(moderate_candidates)} candidate(s) achieved moderate scores and could serve as backup options."
                )
            
            recommendation_parts.append(
                f"We recommend prioritizing <b>{top_candidate['Candidate']}</b> for the next stage of recruitment."
            )
        
        else:
            # Multiple strong candidates
            top_3_strong = strong_candidates.head(3)
            score_range = top_3_strong['Overall'].max() - top_3_strong['Overall'].min()
            
            if score_range < 0.10:  # Very close scores
                names = ", ".join([f"<b>{row['Candidate']}</b>" for _, row in top_3_strong.iterrows()])
                recommendation_parts.append(
                    f"We have <b>{len(strong_candidates)} strong candidates</b> with very similar scores "
                    f"(range: {top_3_strong['Overall'].min():.2f}–{top_3_strong['Overall'].max():.2f}). "
                    f"The top candidates are: {names}."
                )
                recommendation_parts.append(
                    "Given the close scoring, we recommend interviewing multiple candidates to assess "
                    "cultural fit, communication skills, and other qualitative factors."
                )
            else:
                recommendation_parts.append(
                    f"<b>{top_candidate['Candidate']}</b> is the leading candidate with a score of <b>{top_score:.2f}</b>, "
                    f"followed by {len(strong_candidates)-1} other strong candidate(s)."
                )
                recommendation_parts.append(
                    f"We recommend prioritizing <b>{top_candidate['Candidate']}</b>, while keeping other strong candidates "
                    f"as viable alternatives."
                )
        
        # Combine recommendation parts
        for part in recommendation_parts:
            story.append(Paragraph(part, styles['Normal']))
            story.append(Spacer(1, 0.2*cm))
        
        story.append(Spacer(1, 0.3*cm))
        
        # Legend
        story.append(Paragraph("Score Interpretation", heading_style))
        legend_text = (
            "• <b>Strong (≥0.70):</b> Excellent alignment with requirements<br/>"
            "• <b>Moderate (0.45-0.70):</b> Partial or indirect match<br/>"
            "• <b>Weak (&lt;0.45):</b> Likely not covered or weak evidence"
        )
        story.append(Paragraph(legend_text, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        buf.seek(0)
        return buf.getvalue()
    
    except Exception as e:
        st.error(f"PDF generation error: {e}")
        return None

def to_individual_candidate_pdf(
    candidate_name: str,
    coverage_row: pd.Series,
    insights: Dict[str, Any],
    evidence_map: Dict[Tuple[str,str], Tuple[str,float]],
    cat_map: Dict[str, str],
    hi: float = 0.70,
    lo: float = 0.45,
    include_evidence: bool = False,
    job_title: str = ""
) -> Optional[bytes]:
    """Generate individual candidate report PDF."""
    if reportlab is None:
        return None
    
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from datetime import datetime
        
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#1F77B4'),
            spaceAfter=20,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#1F77B4'),
            spaceAfter=10,
            spaceBefore=10
        )
        
        # Title
        story.append(Paragraph(f"Candidate Report: {candidate_name}", title_style))
        if job_title:
            story.append(Paragraph(f"Position: {job_title}", styles['Normal']))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 0.5*cm))
        
        # Overall Score
        overall_score = coverage_row.get('Overall', 0.0)
        rating = "Strong Match" if overall_score >= hi else ("Moderate Match" if overall_score >= lo else "Weak Match")
        
        score_data = [
            ["Overall Score", f"{overall_score:.2f}"],
            ["Rating", rating]
        ]
        score_table = Table(score_data, colWidths=[8*cm, 8*cm])
        score_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#E7F3FF')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(score_table)
        story.append(Spacer(1, 0.5*cm))
        
        # Key Strengths
        story.append(Paragraph("Key Strengths", heading_style))
        strengths = insights.get('top', [])
        if strengths:
            for s in strengths:
                story.append(Paragraph(f"• {s}", styles['Normal']))
        else:
            story.append(Paragraph("<i>No strengths identified</i>", styles['Normal']))
        story.append(Spacer(1, 0.3*cm))
        
        # Development Areas / Gaps
        story.append(Paragraph("Development Areas", heading_style))
        gaps = insights.get('gaps', [])
        if gaps:
            for g in gaps:
                story.append(Paragraph(f"• {g}", styles['Normal']))
        else:
            story.append(Paragraph("<i>No significant gaps identified</i>", styles['Normal']))
        story.append(Spacer(1, 0.3*cm))
        
        # Additional Notes
        notes = insights.get('notes', '')
        if notes:
            story.append(Paragraph("Additional Notes", heading_style))
            story.append(Paragraph(notes, styles['Normal']))
            story.append(Spacer(1, 0.3*cm))
        
        # Detailed Scores by Criteria
        story.append(Paragraph("Detailed Scores by Criteria", heading_style))
        
        # Group by category
        criteria_by_cat = {}
        for col in coverage_row.index:
            if col not in ('Candidate', 'Overall'):
                cat = cat_map.get(col, 'Uncategorized')
                if cat not in criteria_by_cat:
                    criteria_by_cat[cat] = []
                criteria_by_cat[cat].append((col, coverage_row[col]))
        
        for category in sorted(criteria_by_cat.keys()):
            story.append(Paragraph(f"<b>{category}</b>", styles['Heading4']))
            
            criteria_data = [["Criterion", "Score", "Rating"]]
            for crit, score in sorted(criteria_by_cat[category], key=lambda x: x[1], reverse=True):
                rating = "Strong" if score >= hi else ("Moderate" if score >= lo else "Weak")
                # Use Paragraph for criterion name to enable text wrapping
                crit_para = Paragraph(crit, styles['Normal'])
                criteria_data.append([crit_para, f"{score:.2f}", rating])
            
            crit_table = Table(criteria_data, colWidths=[10*cm, 3*cm, 3*cm])
            crit_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E7F3FF')),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
            ]))
            
            # Color code ratings
            for i, (_, score) in enumerate(sorted(criteria_by_cat[category], key=lambda x: x[1], reverse=True), 1):
                if score >= hi:
                    crit_table.setStyle(TableStyle([('BACKGROUND', (2, i), (2, i), colors.HexColor('#C6EFCE'))]))
                elif score >= lo:
                    crit_table.setStyle(TableStyle([('BACKGROUND', (2, i), (2, i), colors.HexColor('#FFEB9C'))]))
                else:
                    crit_table.setStyle(TableStyle([('BACKGROUND', (2, i), (2, i), colors.HexColor('#FFC7CE'))]))
            
            story.append(crit_table)
            story.append(Spacer(1, 0.3*cm))
        
        # Evidence snippets (optional)
        if include_evidence and evidence_map:
            story.append(Paragraph("Evidence Snippets (Selected)", heading_style))
            story.append(Paragraph("<i>Showing evidence for top-scored criteria</i>", styles['Normal']))
            
            # Get top 5 criteria by score
            top_criteria = sorted(
                [(col, coverage_row[col]) for col in coverage_row.index if col not in ('Candidate', 'Overall')],
                key=lambda x: x[1],
                reverse=True
            )[:5]
            
            for crit, score in top_criteria:
                evidence_key = (candidate_name, crit)
                if evidence_key in evidence_map:
                    snippet, _ = evidence_map[evidence_key]
                    story.append(Paragraph(f"<b>{crit}</b> (Score: {score:.2f})", styles['Heading4']))
                    story.append(Paragraph(f"<i>{snippet[:300]}...</i>", styles['Normal']))
                    story.append(Spacer(1, 0.2*cm))
        
        # Build PDF
        doc.build(story)
        buf.seek(0)
        return buf.getvalue()
    
    except Exception as e:
        st.error(f"PDF generation error: {e}")
        return None

# -------- Word/DOCX Export Functions --------

def to_executive_summary_docx(
    coverage: pd.DataFrame,
    insights: Dict[str, Dict],
    jd_text: str,
    cat_map: Dict[str, str],
    hi: float = 0.70,
    lo: float = 0.45,
    jd_filename: str = "Job Description"
) -> Optional[bytes]:
    """Generate editable Word document for executive summary."""
    if docx is None:
        return None
    
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Executive Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('Candidate Analysis Report', 2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Position Section
        doc.add_heading('Position', 1)
        
        # Extract job title
        job_title = "Position Not Specified"
        jd_lines = [line.strip() for line in jd_text.split('\n') if line.strip()]
        
        for line in jd_lines[:5]:
            if len(line) < 10 or 'http' in line.lower() or any(char.isdigit() for char in line[:4]):
                continue
            if any(indicator in line.lower() for indicator in ['position:', 'role:', 'job title:', 'title:']):
                job_title = line.split(':', 1)[1].strip() if ':' in line else line
                break
            elif len(line) > 15 and len(line) < 100:
                job_title = line
                break
        
        p = doc.add_paragraph()
        p.add_run('Position: ').bold = True
        p.add_run(job_title)
        
        # Key Requirements
        jd_summary_lines = []
        title_found = False
        for line in jd_lines:
            if not title_found and job_title in line:
                title_found = True
                continue
            if len(line) > 30 and 'http' not in line.lower() and '@' not in line:
                jd_summary_lines.append(line)
            if len(jd_summary_lines) >= 3:
                break
        
        if jd_summary_lines:
            summary_text = ' '.join(jd_summary_lines)[:400]
            p = doc.add_paragraph()
            p.add_run('Key Requirements: ').bold = True
            p.add_run(f"{summary_text}...")
        
        doc.add_paragraph()  # Spacing
        
        # Analysis Overview
        doc.add_heading('Analysis Overview', 1)
        
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.cell(0, 0).text = 'Metric'
        table.cell(0, 1).text = 'Value'
        table.cell(1, 0).text = 'Total Candidates Analyzed'
        table.cell(1, 1).text = str(len(coverage))
        table.cell(2, 0).text = 'Evaluation Criteria'
        table.cell(2, 1).text = str(len([c for c in coverage.columns if c not in ('Candidate', 'Overall')]))
        table.cell(3, 0).text = 'Analysis Date'
        table.cell(3, 1).text = datetime.now().strftime('%Y-%m-%d')
        
        doc.add_paragraph()  # Spacing
        
        # Top Candidates
        doc.add_heading('Top 5 Candidates', 1)
        
        top5 = coverage[['Candidate', 'Overall']].head(5)
        
        table = doc.add_table(rows=len(top5)+1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Header
        table.cell(0, 0).text = 'Rank'
        table.cell(0, 1).text = 'Candidate Name'
        table.cell(0, 2).text = 'Overall Score'
        table.cell(0, 3).text = 'Rating'
        
        # Data
        for idx, (_, row) in enumerate(top5.iterrows(), 1):
            score = row['Overall']
            rating = "Strong" if score >= hi else ("Moderate" if score >= lo else "Weak")
            
            table.cell(idx, 0).text = str(idx)
            table.cell(idx, 1).text = row['Candidate']
            table.cell(idx, 2).text = f"{score:.2f}"
            table.cell(idx, 3).text = rating
        
        doc.add_paragraph()  # Spacing
        
        # Key Insights for Top 3
        doc.add_heading('Key Insights - Top 3 Candidates', 1)
        
        for idx, (_, row) in enumerate(top5.head(3).iterrows(), 1):
            name = row['Candidate']
            score = row['Overall']
            
            doc.add_heading(f"{idx}. {name} (Score: {score:.2f})", 2)
            
            if name in insights:
                info = insights[name]
                
                # Strengths
                strengths = info.get('top', [])
                if strengths:
                    p = doc.add_paragraph()
                    p.add_run('Strengths:').bold = True
                    for s in strengths[:3]:
                        doc.add_paragraph(s, style='List Bullet')
                
                # Gaps
                gaps = info.get('gaps', [])
                if gaps:
                    p = doc.add_paragraph()
                    p.add_run('Development Areas:').bold = True
                    for g in gaps[:2]:
                        doc.add_paragraph(g, style='List Bullet')
            else:
                doc.add_paragraph('Detailed insights not available', style='Intense Quote')
        
        # Recommendation (same intelligent logic as PDF)
        doc.add_heading('Recommendation', 1)
        
        top_candidate = coverage.iloc[0]
        top_score = top_candidate['Overall']
        
        strong_candidates = coverage[coverage['Overall'] >= hi]
        moderate_candidates = coverage[(coverage['Overall'] >= lo) & (coverage['Overall'] < hi)]
        
        if len(strong_candidates) == 0:
            if len(moderate_candidates) > 0:
                p = doc.add_paragraph()
                run = p.add_run('Caution: ')
                run.bold = True
                p.add_run(f"No candidates achieved a strong match score (≥{hi:.2f}). The highest score was {top_score:.2f} for {top_candidate['Candidate']}, indicating a moderate match.")
                
                doc.add_paragraph(f"We recommend carefully reviewing the {len(moderate_candidates)} moderate-scoring candidate(s) to identify specific skill gaps, or consider expanding the candidate pool.")
            else:
                p = doc.add_paragraph()
                run = p.add_run('Warning: ')
                run.bold = True
                p.add_run(f"All candidates scored below the moderate threshold ({lo:.2f}). The highest score was only {top_score:.2f} for {top_candidate['Candidate']}.")
                
                doc.add_paragraph("We recommend reconsidering the job requirements or sourcing additional candidates, as the current pool shows weak alignment with the position criteria.")
        
        elif len(strong_candidates) == 1:
            doc.add_paragraph(f"{top_candidate['Candidate']} is the clear leading candidate with a strong overall score of {top_score:.2f}, demonstrating excellent alignment with the position requirements.")
            
            if len(moderate_candidates) > 0:
                doc.add_paragraph(f"Additionally, {len(moderate_candidates)} candidate(s) achieved moderate scores and could serve as backup options.")
            
            p = doc.add_paragraph(f"We recommend prioritizing {top_candidate['Candidate']} for the next stage of recruitment.")
            p.runs[0].bold = True
        
        else:
            top_3_strong = strong_candidates.head(3)
            score_range = top_3_strong['Overall'].max() - top_3_strong['Overall'].min()
            
            if score_range < 0.10:
                names = ", ".join([row['Candidate'] for _, row in top_3_strong.iterrows()])
                doc.add_paragraph(f"We have {len(strong_candidates)} strong candidates with very similar scores (range: {top_3_strong['Overall'].min():.2f}–{top_3_strong['Overall'].max():.2f}). The top candidates are: {names}.")
                doc.add_paragraph("Given the close scoring, we recommend interviewing multiple candidates to assess cultural fit, communication skills, and other qualitative factors.")
            else:
                doc.add_paragraph(f"{top_candidate['Candidate']} is the leading candidate with a score of {top_score:.2f}, followed by {len(strong_candidates)-1} other strong candidate(s).")
                doc.add_paragraph(f"We recommend prioritizing {top_candidate['Candidate']}, while keeping other strong candidates as viable alternatives.")
        
        # Legend
        doc.add_heading('Score Interpretation', 1)
        doc.add_paragraph(f"Strong (≥{hi:.2f}): Excellent alignment with requirements", style='List Bullet')
        doc.add_paragraph(f"Moderate ({lo:.2f}–{hi:.2f}): Partial or indirect match", style='List Bullet')
        doc.add_paragraph(f"Weak (<{lo:.2f}): Likely not covered or weak evidence", style='List Bullet')
        
        # Save to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    
    except Exception as e:
        st.error(f"Word document generation error: {e}")
        return None

def to_individual_candidate_docx(
    candidate_name: str,
    coverage_row: pd.Series,
    insights: Dict[str, Any],
    cat_map: Dict[str, str],
    hi: float = 0.70,
    lo: float = 0.45,
    job_title: str = ""
) -> Optional[bytes]:
    """Generate editable Word document for individual candidate."""
    if docx is None:
        return None
    
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f'Candidate Report: {candidate_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if job_title:
            position_para = doc.add_paragraph(f"Position: {job_title}")
            position_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Overall Score
        overall_score = coverage_row.get('Overall', 0.0)
        rating = "Strong Match" if overall_score >= hi else ("Moderate Match" if overall_score >= lo else "Weak Match")
        
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Light Grid Accent 1'
        table.cell(0, 0).text = 'Overall Score'
        table.cell(0, 1).text = f"{overall_score:.2f}"
        table.cell(1, 0).text = 'Rating'
        table.cell(1, 1).text = rating
        
        doc.add_paragraph()  # Spacing
        
        # Key Strengths
        doc.add_heading('Key Strengths', 1)
        strengths = insights.get('top', [])
        if strengths:
            for s in strengths:
                doc.add_paragraph(s, style='List Bullet')
        else:
            doc.add_paragraph('No strengths identified', style='Intense Quote')
        
        # Development Areas
        doc.add_heading('Development Areas', 1)
        gaps = insights.get('gaps', [])
        if gaps:
            for g in gaps:
                doc.add_paragraph(g, style='List Bullet')
        else:
            doc.add_paragraph('No significant gaps identified', style='Intense Quote')
        
        # Additional Notes
        notes = insights.get('notes', '')
        if notes:
            doc.add_heading('Additional Notes', 1)
            doc.add_paragraph(notes)
        
        # Detailed Scores by Criteria
        doc.add_heading('Detailed Scores by Criteria', 1)
        
        # Group by category
        criteria_by_cat = {}
        for col in coverage_row.index:
            if col not in ('Candidate', 'Overall'):
                cat = cat_map.get(col, 'Uncategorized')
                if cat not in criteria_by_cat:
                    criteria_by_cat[cat] = []
                criteria_by_cat[cat].append((col, coverage_row[col]))
        
        for category in sorted(criteria_by_cat.keys()):
            doc.add_heading(category, 2)
            
            sorted_criteria = sorted(criteria_by_cat[category], key=lambda x: x[1], reverse=True)
            
            table = doc.add_table(rows=len(sorted_criteria)+1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Header
            table.cell(0, 0).text = 'Criterion'
            table.cell(0, 1).text = 'Score'
            table.cell(0, 2).text = 'Rating'
            
            # Data
            for idx, (crit, score) in enumerate(sorted_criteria, 1):
                rating = "Strong" if score >= hi else ("Moderate" if score >= lo else "Weak")
                table.cell(idx, 0).text = crit
                table.cell(idx, 1).text = f"{score:.2f}"
                table.cell(idx, 2).text = rating
        
        # Save to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    
    except Exception as e:
        st.error(f"Word document generation error: {e}")
        return None

# -------- GPT Insights helpers --------

def _redact(s: str) -> str:
    """Redact emails, phones, URLs. Keeps text readable."""
    if not s:
        return s
    s = re.sub(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", "[-email]", s)
    s = re.sub(r"\b(?:\+?\d[\d \-\(\)]{8,}\d)\b", "[-phone]", s)
    s = re.sub(r"https?://\S+|www\.\S+", "[-url]", s)
    return s

def _model_for_engine(engine: str) -> Optional[str]:
    if engine == "gpt-4o":
        return "gpt-4o"
    if engine == "gpt-4o":
        return "gpt-4o"
    return None

def gpt_candidate_insights(
    candidate_name: str,
    candidate_text: str,
    jd_text: str,
    criteria: List[str],
    coverage_row: Dict[str, float],
    evidence_map: Dict[Tuple[str,str], Tuple[str,float]]
) -> Dict[str, Any]:
    """
    Returns dict with keys: {'top': [...], 'gaps': [...], 'notes': str}
    Uses JSON schema to keep output consistent. Always uses GPT-4o.
    """
    model = "gpt-4o"

    # Build evidence lines from coverage scores and evidence_map
    ev_items = []
    for c in criteria:
        sc = float(coverage_row.get(c, 0.0))
        snip = evidence_map.get((candidate_name, c), ("", sc))[0]
        _snip_clean = (snip or "")[:400].replace("\n", " ")
        ev_items.append(f"- {c} (score {sc:.2f}): {_snip_clean}")
    ev_lines = ev_items

    cand_body = candidate_text
    jd_body = jd_text

    system = (
        "You are a hiring analyst. Produce concise, high-signal insights about a candidate "
        "relative to the provided Job Description and evidence. Be specific and avoid fluff."
    )
    schema = {
        "type": "object",
        "properties": {
            "top": {"type": "array", "items": {"type": "string"}},
            "gaps": {"type": "array", "items": {"type": "string"}},
            "notes": {"type": "string"},
        },
        "required": ["top", "gaps", "notes"],
        "additionalProperties": False,
    }
    user = (
        f"Job Description (excerpted):\n{jd_body[:3000]}\n\n"
        f"Candidate: {candidate_name}\n"
        f"Candidate content (excerpted):\n{cand_body[:3000]}\n\n"
        "Evidence by criterion (top 10 by similarity):\n"
        + "\n".join(ev_lines) +
        "\n\nTask:\n"
        "- Provide 3–6 bullet **Top strengths** tied to criteria and tangible evidence.\n"
        "- Provide 3–6 bullet **Gaps / risks** with rationale.\n"
        "- Provide a short **Notes** paragraph (2–4 sentences) with an overall view."
    )

    try:
        data = call_llm_json(system, user, schema, model=model)
        # Normalize
        return {
            "top": [normalize_ws(x) for x in data.get("top", []) if normalize_ws(x)],
            "gaps": [normalize_ws(x) for x in data.get("gaps", []) if normalize_ws(x)],
            "notes": normalize_ws(data.get("notes", "")),
        }
    except Exception:
        return {}

def _update_stage(msg: str):
    """Placeholder for status updates during analysis"""
    pass



with st.sidebar:
    st.markdown("## 🧭 Navigation")
    
    # Initialize page if not set
    if "current_page" not in st.session_state:
        st.session_state.current_page = "Upload & Analyse"
    
    # Radio navigation - better use of vertical space
    pages = ["Upload & Analyse", "Job Criteria", "Scoring Analysis", "Candidate Insights", "Export Reports", "Settings", "Help & FAQ"]
    
    selected_page = st.radio(
        "Navigation",
        pages,
        key="nav_radio",
        label_visibility="collapsed"
    )
    
    # Only update state if changed
    if selected_page != st.session_state.current_page:
        st.session_state.current_page = selected_page
    
    # ===== Cost Estimate (if data available) =====
    current_jd = st.session_state.get("jd")
    current_candidates = st.session_state.get("cached_candidates", [])
    
    if current_jd and current_candidates:
        st.markdown("---")
        st.markdown("### 💰 Analysis Cost")
        
        # Get criteria count
        criteria_text = st.session_state.get("criteria_text", "")
        if criteria_text:
            num_criteria = len([c for c in criteria_text.splitlines() if c.strip()])
        else:
            num_criteria = 20  # estimate
        
        # Calculate cost
        jd_text = current_jd.text
        candidate_texts = [c.text for c in current_candidates]
        cost_info = estimate_analysis_cost(jd_text, candidate_texts, num_criteria)
        
        # Display compact metrics
        st.metric("Estimated Total", f"${cost_info['estimated_cost_usd']:.2f}")
        st.caption(f"${cost_info['cost_per_candidate']:.2f} per candidate")
        st.caption(f"{len(current_candidates)} candidate{'s' if len(current_candidates) != 1 else ''}")
        
        # Expandable breakdown
        with st.expander("💡 Details"):
            st.write(f"**JD Extraction:** ${cost_info['breakdown']['jd_extraction']:.2f}")
            st.write(f"**Insights:** ${cost_info['breakdown']['candidate_insights']:.2f}")
            st.caption(f"Model: {cost_info['model']}")
            st.caption("Estimate may vary ±10% from actual usage")

# -----------------------------
# File Processing (needs to happen before page rendering)
# -----------------------------
# Initialize variables
jd_file = None
jd_manual = ""
uploaded = []

# Only process uploads on the Upload & Analyse page to keep them in scope
# But we need to maintain state across page navigation

# Only process uploads on the Upload & Analyse page to keep them in scope
# But we need to maintain state across page navigation

# -----------------------------
# Page Rendering Based on Selection
# -----------------------------
current_page = st.session_state.get("current_page", "Upload & Analyse")

if current_page == "Upload & Analyse":
    # Status container for analysis progress (at the very top)
    st.session_state._status_container = st.empty()
    
    # Add spacing below header
    st.markdown("<div style='margin-top: 1.5rem;'></div>", unsafe_allow_html=True)
    
    # Simplified, cleaner layout
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown("#### 📄 Job Description (JD)")
        jd_file = st.file_uploader("Upload JD", type=["pdf","docx","txt"], key="jd_upl", accept_multiple_files=False, label_visibility="collapsed")
        
        with st.expander("✏️ Paste JD text manually"):
            jd_manual = st.text_area("JD text", value=st.session_state.get("jd_manual_text",""), height=180, label_visibility="collapsed")
            st.session_state.jd_manual_text = jd_manual
    
    with col2:
        st.markdown("#### 👥 Candidate Resumes")
        uploaded = st.file_uploader("Upload resumes", accept_multiple_files=True, type=["pdf","docx","txt"], key="cand_upl", label_visibility="collapsed")
    
    # Show currently loaded files with option to remove
    current_jd = st.session_state.get("jd")
    current_candidates = st.session_state.get("cached_candidates", [])
    
    if current_jd or current_candidates:
        with st.expander("📋 Currently Loaded Files", expanded=False):
            if current_jd:
                st.markdown("**Job Description:**")
                col_jd, col_jd_btn = st.columns([4, 1])
                with col_jd:
                    st.text(current_jd.file_name)
                with col_jd_btn:
                    if st.button("❌", key="remove_jd", help="Remove JD"):
                        st.session_state.jd = None
                        st.session_state.last_coverage = pd.DataFrame()
                        st.session_state.last_insights = {}
                        st.rerun()
            
            if current_candidates:
                st.markdown("**Candidates:**")
                for idx, cand in enumerate(current_candidates):
                    col_name, col_file, col_btn = st.columns([2, 2, 1])
                    with col_name:
                        st.text(cand.name)
                    with col_file:
                        st.caption(cand.file_name)
                    with col_btn:
                        if st.button("❌", key=f"remove_cand_{idx}", help=f"Remove {cand.name}"):
                            # Remove this candidate from the list
                            updated_candidates = [c for c in current_candidates if c.hash != cand.hash]
                            st.session_state.cached_candidates = updated_candidates
                            # Clear analysis results since data changed
                            st.session_state.last_coverage = pd.DataFrame()
                            st.session_state.last_insights = {}
                            st.rerun()
    
    # Analyse button with spacing
    st.markdown("<div style='margin-top: 1rem;'></div>", unsafe_allow_html=True)
    if st.button("🔍 Analyse / Update", type="primary", width='stretch'):
        st.session_state._trigger_analyse = True
        st.rerun()
    
    # JD assignment from file or manual text
    if jd_file is not None:
        b = jd_file.read()
        if jd_file.name.lower().endswith('.pdf'):
            if st.session_state.get("use_layout_pdf", True):
                try:
                    jd_text = extract_text_layout_aware(b, jd_file.name)
                    if not jd_text.strip():
                        jd_text = read_file_bytes(b, jd_file.name)
                except Exception:
                    jd_text = read_file_bytes(b, jd_file.name)
            else:
                jd_text = read_file_bytes(b, jd_file.name)
        else:
            jd_text = read_file_bytes(b, jd_file.name)
        st.session_state.jd = JD(file_name=jd_file.name, text=jd_text, hash=hash_bytes(b), raw_bytes=b)
    elif jd_manual.strip():
        st.session_state.jd = JD(file_name="Manual JD", text=jd_manual, hash=hash_text(jd_manual), raw_bytes=None)

    # Candidates ingest
    new_candidates: List[Candidate] = []
    if uploaded:
        for uf in uploaded:
            b = uf.read()
            if uf.name.lower().endswith('.pdf'):
                if st.session_state.get("use_layout_pdf", True):
                    try:
                        text = extract_text_layout_aware(b, uf.name)
                        if not text.strip():
                            text = read_file_bytes(b, uf.name)
                    except Exception:
                        text = read_file_bytes(b, uf.name)
                else:
                    text = read_file_bytes(b, uf.name)
            else:
                text = read_file_bytes(b, uf.name)
            name_guess = infer_candidate_name(uf.name, text)
            new_candidates.append(Candidate(name=name_guess, file_name=uf.name, text=text, hash=hash_bytes(b), raw_bytes=b))
    if new_candidates:
        existing = {c.hash: c for c in st.session_state.get("cached_candidates", [])}
        for c in new_candidates: existing[c.hash] = c
        st.session_state.cached_candidates = list(existing.values())
    
    # Overview/Status Section with better visual hierarchy
    st.markdown("<div style='margin-top: 1.5rem;'></div>", unsafe_allow_html=True)
    
    # Custom styling for smaller metrics
    st.markdown("""
        <style>
        [data-testid="stMetricValue"] {
            font-size: 1.2rem;
        }
        [data-testid="stMetricLabel"] {
            font-size: 0.9rem;
        }
        </style>
    """, unsafe_allow_html=True)
    
    col_a, col_b, col_c = st.columns(3)
    
    with col_a:
        jd_status = "📄 Loaded" if st.session_state.get("jd") else "⚠️ Not loaded"
        st.metric("Job Description", jd_status)
        if st.session_state.get("jd"):
            st.caption(f"{len(st.session_state.jd.text):,} characters")
    
    with col_b:
        num_cands = len(st.session_state.get("cached_candidates", []))
        st.metric("Candidates", num_cands)
    
    with col_c:
        covdf = st.session_state.get("last_coverage", pd.DataFrame())
        num_criteria = len([c for c in covdf.columns if c not in ('Candidate','Overall')]) if isinstance(covdf, pd.DataFrame) and not covdf.empty else 0
        st.metric("Criteria", num_criteria)
    
    # Show top results if analysis has been run
    if isinstance(covdf, pd.DataFrame) and not covdf.empty:
        # Subtle info text explaining scores - full width above results
        st.markdown(
            "<div style='padding: 0.5rem 0; color: #666; font-size: 0.9rem; margin-bottom: 0.5rem;'>"
            "Scores indicate the strength of the match between candidate resumes and the JD criteria. 0.00 = no match, 1.00 = perfect match.<br>"
            "Icons indicate score strength in 3 ranges: ✅ Strong (≥0.70) | ⚠️ Moderate (0.45–0.70) | ⛔ Weak (&lt;0.45)"
            "</div>",
            unsafe_allow_html=True
        )
        
        result_col1, result_col2 = st.columns([1, 1])
        
        with result_col1:
            st.markdown("### 🏆 Top Candidates")
            try:
                top3 = covdf[["Candidate","Overall"]].head(3)
                hi = st.session_state.get("cov_hi", 0.70)
                lo = st.session_state.get("cov_lo", 0.45)
                for idx, r in top3.iterrows():
                    score = r['Overall']
                    icon = "✅" if score >= hi else ("⚠️" if score >= lo else "⛔")
                    st.write(f"**{idx+1}.** {r['Candidate']} — score: {score:.2f} {icon}")
            except Exception:
                pass
        
        with result_col2:
            st.markdown("### 📈 Top Matched Criteria")
            try:
                # Calculate average scores by criterion
                crit_cols = [c for c in covdf.columns if c not in ('Candidate','Overall')]
                avg_by_criterion = {}
                for crit in crit_cols:
                    scores = covdf[crit].tolist()
                    avg_by_criterion[crit] = sum(scores) / len(scores) if scores else 0
                
                # Display top 5 criteria
                sorted_crit = sorted(avg_by_criterion.items(), key=lambda x: x[1], reverse=True)
                top_5 = sorted_crit[:5]
                
                hi = st.session_state.get("cov_hi", 0.70)
                lo = st.session_state.get("cov_lo", 0.45)
                for crit, score in top_5:
                    icon = "✅" if score >= hi else ("⚠️" if score >= lo else "⛔")
                    st.write(f"• {crit}: {score:.2f} {icon}")
            except Exception:
                pass

elif current_page == "Job Criteria":
    st.markdown("<div style='margin-top: 1.5rem;'></div>", unsafe_allow_html=True)
    st.markdown("<h2 style='margin-bottom: 0;'>Job Criteria <span style='color: #999; font-weight: normal;'>- Extracted from Job Description (JD)</span></h2>", unsafe_allow_html=True)
    
    # Check if JD exists
    if not st.session_state.get("jd"):
        st.warning("Upload a JD file or paste JD text on the **Upload & Analyse** tab to get started.")
        st.stop()
    
    jd_obj = st.session_state.jd
    
    # Auto-extract sections and criteria (GPT-only) for current JD if not already present
    if API_KEY_SET:
        need_extract = (
            "extracted_hash" not in st.session_state
            or st.session_state.get("extracted_hash") != jd_obj.hash
        )
        if need_extract:
            secs = cached_extract_jd_sections_with_gpt(jd_obj.hash, jd_obj.text)
            st.session_state["extracted_sections"] = _sections_to_dict(secs)
            crits, cat_map = build_criteria_from_gpt_sections(secs, per_section=999, cap_total=10000)
            st.session_state["cat_map"] = cat_map
            st.session_state["criteria_text"] = "\n".join(crits)
            st.session_state["extracted_hash"] = jd_obj.hash
    
    # Compact, collapsible section with all JD details
    with st.expander("📋 Job Description & Extraction Details", expanded=False):
        # Summary metrics in compact row
        col1, col2, col3, col4 = st.columns(4)
        num_criteria = len([l for l in (st.session_state.get("criteria_text","") or "").splitlines() if l.strip()])
        cat_map = st.session_state.get("cat_map", {})
        num_categories = len(set(cat_map.values())) if cat_map else 0
        
        with col1:
            st.metric("JD File", jd_obj.file_name.split('.')[0][:12] + "..." if len(jd_obj.file_name) > 15 else jd_obj.file_name)
        with col2:
            st.metric("Characters", f"{len(jd_obj.text):,}")
        with col3:
            st.metric("Criteria", num_criteria)
        with col4:
            st.metric("Categories", num_categories if num_categories > 0 else "Uncategorized")
        
        st.markdown("---")
        
        # View original file inline (in collapsible section)
        if jd_obj.raw_bytes:
            if jd_obj.file_name.lower().endswith('.pdf'):
                # PDF viewer in expander
                with st.expander("📄 View Original PDF File", expanded=False):
                    import base64
                    base64_pdf = base64.b64encode(jd_obj.raw_bytes).decode('utf-8')
                    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600" type="application/pdf"></iframe>'
                    st.markdown(pdf_display, unsafe_allow_html=True)
                    st.caption("💡 Use the browser's PDF controls to zoom, navigate pages, or download if needed.")
            elif jd_obj.file_name.lower().endswith(('.docx', '.doc')):
                with st.expander("📝 Original DOCX/DOC File", expanded=False):
                    st.info("DOCX/DOC files cannot be previewed inline. View the 'Full Text Extract' below, or download the file to open in Word.", icon="ℹ️")
                    # Provide download button for DOCX
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if jd_obj.file_name.lower().endswith('.docx') else "application/msword"
                    st.download_button(
                        "📥 Download to View in Word",
                        data=jd_obj.raw_bytes,
                        file_name=jd_obj.file_name,
                        mime=mime_type,
                        help="Download to view in Microsoft Word"
                    )
            else:
                # Plain text files in expander
                with st.expander("📄 View Original Text File", expanded=False):
                    st.code(jd_obj.text, language=None, line_numbers=False)
        
        st.markdown("---")
        
        # Extracted sections (if any)
        secs_dict = st.session_state.get("extracted_sections", None)
        if secs_dict:
            st.markdown("**📋 Extracted Sections (auto from GPT)**")
            sec_col1, sec_col2 = st.columns(2)
            with sec_col1:
                st.markdown("**🎯 Key skills**")
                for x in secs_dict.get("key_skills", []): st.markdown(f"- {x}")
                st.markdown("**📌 Responsibilities**")
                for x in secs_dict.get("responsibilities", []): st.markdown(f"- {x}")
            with sec_col2:
                st.markdown("**🎓 Qualifications**")
                for x in secs_dict.get("qualifications", []): st.markdown(f"- {x}")
                st.markdown("**💼 Experience required**")
                for x in secs_dict.get("experience_required", []): st.markdown(f"- {x}")
            
            st.markdown("---")
        elif not API_KEY_SET:
            st.info("💡 Set `OPENAI_API_KEY` in your environment to auto-extract sections and build criteria.", icon="ℹ️")
            st.markdown("---")
        
        # Full text extract
        st.markdown("**📄 Full Text Extract** - used for AI/GPT analysis")
        st.text_area("Full text extracted from JD", value=jd_obj.text, height=200, label_visibility="collapsed", key="jd_text_preview")

    # --- Review & edit criteria ---
    st.markdown("### ✏️ Review & Edit Criteria")
    
    init_lines = [l for l in (st.session_state.get("criteria_text","") or "").splitlines() if l.strip()]
    cat_map = st.session_state.get("cat_map", {})
    
    # Show count and category breakdown
    num_criteria = len(init_lines)
    categories = set(cat_map.get(c, "Uncategorized") for c in init_lines)
    
    st.caption(f"Editing **{num_criteria}** criteria across **{len(categories)}** categories")
    
    st.info("""💡 **Tip:** Add, remove, or reword criteria. Uncheck **Use** to exclude items. Changes take effect after clicking "Save & Re-run Analysis".""", icon="ℹ️")

    # Prepare editable table (Use, Category, Criterion)
    import pandas as pd
    rows = []
    for c in init_lines:
        rows.append({"Use": True, "Category": cat_map.get(c, ""), "Criterion": c})
    if not rows:
        rows = [{"Use": True, "Category": "", "Criterion": ""}]
    df = pd.DataFrame(rows, columns=["Use","Category","Criterion"])

    edited = st.data_editor(
        df,
        width='stretch',
        num_rows="dynamic",
        column_config={
            "Use": st.column_config.CheckboxColumn("Use", help="Uncheck to exclude from analysis."),
            "Category": st.column_config.TextColumn("Category", help="Optional grouping label for display order."),
            "Criterion": st.column_config.TextColumn("Criterion", help="Edit the text of the criterion."),
        },
        hide_index=True,
    )

    
    # --- Action buttons in card-style layout ---
    st.markdown("#### ⚙️ Actions")

    action_col1, action_col2, action_col3 = st.columns(3)

    with action_col1:
        st.markdown("**📤 Export Criteria**")
        st.caption("Save selected criteria to CSV")
        # Export selected rows as CSV with Criterion,Category
        import pandas as _pd
        try:
            _rows = [
                {"Criterion": str(c).strip(),
                "Category": (str(cat).strip() if str(cat).strip() else "")}
                for u, c, cat in zip(
                    edited["Use"],
                    edited["Criterion"],
                    edited.get("Category", [""] * len(edited)),
                )
                if u and str(c).strip()
            ]
        except Exception:
            _rows = [
                {"Criterion": ln, "Category": ""}
                for ln in (st.session_state.get("criteria_text", "") or "").splitlines()
                if ln.strip()
            ]
        _csv = _pd.DataFrame(_rows, columns=["Criterion", "Category"]).to_csv(index=False)
        st.download_button(
            "Export to CSV",
            data=_csv,
            file_name="criteria.csv",
            mime="text/csv",
            use_container_width=True
        )

    with action_col2:
        st.markdown("**📥 Import Criteria**")
        st.caption("Load from CSV file")
        _uploaded = st.file_uploader(
            "Upload CSV", type=["csv"], accept_multiple_files=False, key="crit_csv_upl", label_visibility="collapsed"
        )
        if st.button("Import & Replace", key="import_replace_btn", use_container_width=True):
            import pandas as _pd
            if _uploaded is None:
                st.warning("Upload a CSV file first.")
            else:
                try:
                    _df = _pd.read_csv(_uploaded)

                    # Normalize expected columns
                    rename_map = {}
                    for cn in _df.columns:
                        cl = str(cn).lower().strip()
                        if cl in ("criterion", "criteria"):
                            rename_map[cn] = "Criterion"
                        elif cl in ("category", "cat"):
                            rename_map[cn] = "Category"
                    if rename_map:
                        _df = _df.rename(columns=rename_map)

                    if "Criterion" not in _df.columns:
                        st.error("CSV must include a 'Criterion' column.")
                    else:
                        _df["Criterion"] = (
                            _df["Criterion"].fillna("").astype(str).str.strip()
                        )
                        if "Category" not in _df.columns:
                            _df["Category"] = ""
                        _df["Category"] = (
                            _df["Category"].fillna("").astype(str).str.strip()
                        )
                        _df = _df[_df["Criterion"] != ""]
                        crit_list = _df["Criterion"].tolist()
                        cat_map = {
                            r["Criterion"]: r["Category"]
                            for _, r in _df.iterrows()
                            if str(r["Category"]).strip()
                        }
                        if crit_list:
                            st.session_state["criteria_text"] = "\n".join(crit_list)
                            st.session_state["cat_map"] = cat_map
                            st.session_state["_trigger_analyse"] = True
                            st.rerun()
                        else:
                            st.warning("No valid criteria found in CSV.")
                except Exception as e:
                    st.error(f"Failed to read CSV: {e}")

    with action_col3:
        st.markdown("**🔄 Reset Criteria**")
        st.caption("Restore to auto-generated from JD")
        if st.button("Reset to Auto-Generated", use_container_width=True):
            if st.session_state.get("jd") and API_KEY_SET:
                _jd = st.session_state.jd
                _secs = cached_extract_jd_sections_with_gpt(_jd.hash, _jd.text)
                _crits, _cat_map = build_criteria_from_gpt_sections(_secs, per_section=999, cap_total=10000)
                st.session_state["extracted_sections"] = _sections_to_dict(_secs)
                st.session_state["criteria_text"] = "\n".join(_crits)
                st.session_state["cat_map"] = _cat_map
                st.session_state["_trigger_analyse"] = True
                st.rerun()
            else:
                st.warning("Cannot reset without a JD and API key.")

 
    # Primary save button at bottom
    save_col1, save_col2 = st.columns([1, 3])
    with save_col1:
        if st.button("💾 Save & Re-run Analysis", type="primary", use_container_width=True):
            # Persist edited criteria & categories
            new_use = list(edited["Use"])
            new_crits = list(edited["Criterion"])
            new_cats = list(edited["Category"]) if "Category" in edited else [""] * len(new_crits)

            kept = []
            new_cat_map = {}
            for used, crit, cat in zip(new_use, new_crits, new_cats):
                crit_s = str(crit).strip()
                if used and crit_s:
                    kept.append(crit_s)
                    if cat and str(cat).strip():
                        new_cat_map[crit_s] = str(cat).strip()
            st.session_state["criteria_text"] = "\n".join(kept)
            # Merge existing cat_map to retain known mappings unless overridden
            old_map = st.session_state.get("cat_map", {})
            old_map.update(new_cat_map)
            st.session_state["cat_map"] = {k:v for k,v in old_map.items() if k in kept}

            # Trigger fresh analysis
            st.session_state["_trigger_analyse"] = True
            st.rerun()
    
    with save_col2:
        pass

elif current_page == "Scoring Analysis":
    st.markdown("<div style='margin-top: 1.5rem;'></div>", unsafe_allow_html=True)
    
    coverage_df = st.session_state.get("last_coverage", pd.DataFrame())
    
    # Header row with summary and top candidates in corner
    header_col1, header_col2 = st.columns([2, 1])
    
    with header_col1:
        st.subheader("Scoring Analysis")
        if not isinstance(coverage_df, pd.DataFrame) or coverage_df.empty:
            st.warning("Upload JD/criteria and candidate resumes, then click Analyse.")
        else:
            # Compact summary line
            num_candidates = len(coverage_df)
            num_criteria = len([c for c in coverage_df.columns if c not in ("Candidate","Overall")])
            st.markdown(f"Analyzing **{num_candidates} candidates** across **{num_criteria} criteria**")
    
    with header_col2:
        # Compact top candidates list in corner (only if we have data)
        # Will be populated after checkbox is rendered
        top_candidates_placeholder = st.empty()
    
    # Only continue if we have data
    if not isinstance(coverage_df, pd.DataFrame) or coverage_df.empty:
        st.stop()
    
    # Info box with scoring guide - full width, below header, smaller font
    st.markdown(
        "<div style='font-size: 0.9rem;'>"
        "<div style='padding: 0.75rem; background-color: #e7f3ff; border-left: 4px solid #1f77b4; border-radius: 0.25rem; margin-bottom: 1rem;'>"
        """Each score (0.00–1.00) measures how well a candidate's resume matches a criterion. Use "Scores" tickbox to show/hide numeric scores.<br>"""
        "Icons indicate score strength in 3 ranges: ✅ Strong (≥0.70) | ⚠️ Moderate (0.45–0.70) | ⛔ Weak (&lt;0.45)"
        "</div></div>",
        unsafe_allow_html=True
    )
    
    # Get threshold settings (used throughout this page)
    hi = st.session_state.get("cov_hi", 0.70)
    lo = st.session_state.get("cov_lo", 0.45)
    
    # Compact single-row controls (no header needed - matrix is obvious)
    st.markdown("""
        <style>
        .compact-controls {
            font-size: 0.85rem;
            color: #666;
            margin: 0.5rem 0;
        }
        /* Make controls more compact */
        div[data-testid="stHorizontalBlock"] > div {
            gap: 0.5rem !important;
        }
        </style>
    """, unsafe_allow_html=True)
    
    ctrl_col1, ctrl_col2, ctrl_col3, ctrl_col4, ctrl_col5 = st.columns([2, 1.5, 1.5, 1, 0.8])
    
    with ctrl_col1:
        sort_options = ["Name (A-Z)", "Overall Score (High-Low)", "Overall Score (Low-High)"]
        # Reset index if it's out of range (e.g., from previous session with more options)
        default_idx = st.session_state.get("cov_sort_idx", 1)
        if default_idx >= len(sort_options):
            default_idx = 1
        sort_by = st.selectbox("Sort candidates by", 
                              sort_options,
                              index=default_idx,
                              key="sort_select",
                              label_visibility="visible")
        st.session_state.cov_sort_idx = sort_options.index(sort_by)
    
    with ctrl_col2:
        filter_options = ["All", "Top 3", "Top 5", "Top 10", "Custom Selection"]
        show_top_n = st.selectbox("Show", 
                                 filter_options,
                                 index=0,
                                 key="filter_select",
                                 label_visibility="visible")
    
    # Custom candidate selection (only shown if "Custom Selection" is chosen)
    selected_candidates = None
    if show_top_n == "Custom Selection":
        all_candidates = coverage_df["Candidate"].tolist()
        st.markdown("**👥 Select candidates to compare:**")
        selected_candidates = st.multiselect(
            "Select candidates",
            all_candidates,
            default=all_candidates[:3] if len(all_candidates) >= 3 else all_candidates,
            key="custom_candidate_selector",
            label_visibility="collapsed"
        )
        if not selected_candidates:
            st.warning("⚠️ Select at least one candidate to display.")
    
    with ctrl_col3:
        expand_categories = st.checkbox("Expand categories", value=st.session_state.get("cov_expand_all", True))
        st.session_state.cov_expand_all = expand_categories
    
    with ctrl_col4:
        # Checkbox that syncs with session state - let Streamlit manage the state naturally
        show_scores_checked = st.checkbox(
            "Scores", 
            value=st.session_state.get("cov_show_scores", False),
            help="Show numeric scores",
            key="cov_show_scores_checkbox"
        )
        # Update session state for next render
        st.session_state.cov_show_scores = show_scores_checked
        # Use the checkbox value directly (not session state)
        show_scores = show_scores_checked
    
    with ctrl_col5:
        # CSV download button placeholder - will be populated after sorting
        csv_placeholder = st.empty()
    
    # Now render top candidates with the current show_scores value
    with top_candidates_placeholder.container():
        if isinstance(coverage_df, pd.DataFrame) and not coverage_df.empty:
            st.markdown("##### 🏆 Top Candidates")
            
            # Sort for top candidates display
            top_5_df = coverage_df.sort_values("Overall", ascending=False).head(5)
            
            # Display as simple vertical list using current show_scores value
            for idx, row in top_5_df.iterrows():
                score = row["Overall"]
                icon = "✅" if score >= hi else ("⚠️" if score >= lo else "⛔")
                # Get position (1-5)
                pos = top_5_df.index.get_loc(idx) + 1
                if show_scores:
                    st.markdown(f"**{pos}.** {row['Candidate']} — {score:.2f} {icon}")
                else:
                    st.markdown(f"**{pos}.** {row['Candidate']} — {icon}")

    # Sort candidates based on selection
    if sort_by == "Overall Score (High-Low)":
        sorted_coverage = coverage_df.sort_values("Overall", ascending=False)
    elif sort_by == "Overall Score (Low-High)":
        sorted_coverage = coverage_df.sort_values("Overall", ascending=True)
    else:  # Name (A-Z)
        sorted_coverage = coverage_df.sort_values("Candidate")
    
    # Filter based on selection type
    if show_top_n == "Custom Selection":
        # Use custom multiselect
        if selected_candidates:
            sorted_coverage = sorted_coverage[sorted_coverage["Candidate"].isin(selected_candidates)]
        else:
            # No candidates selected - show nothing (warning already displayed above)
            sorted_coverage = pd.DataFrame()
    elif show_top_n != "All":
        # Use Top N filter
        n = int(show_top_n.split()[-1])
        sorted_coverage = sorted_coverage.head(n)
    
    # Stop here if no candidates to display (must check before accessing columns)
    if sorted_coverage.empty:
        st.stop()
    
    candidates = sorted_coverage["Candidate"].tolist()
    
    # Now populate the CSV download button in the placeholder
    with csv_placeholder:
        csv_data = sorted_coverage.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="📥",
            data=csv_data,
            file_name="coverage_matrix.csv",
            mime="text/csv",
            help="Download CSV",
            key="csv_download_btn"
            )

    # Hint text for sorting
    st.markdown(
        "<p style='font-size: 0.85rem; color: #666; margin: 0.5rem 0 1rem 0;'>"
        "💡 Hint: If you want to see which Criteria achieved the best (or worst) average scores, "
        "just click on the column header \"Average\" to sort by that column."
        "</p>",
        unsafe_allow_html=True
    )
    
    # Build coverage matrix with sorted/filtered candidates
    cat_map = st.session_state.get("cat_map", {})
    crit_cols = [c for c in sorted_coverage.columns if c not in ("Candidate","Overall")]
    mat = sorted_coverage.set_index("Candidate")[crit_cols].T
    ordered = sorted(crit_cols, key=lambda c: (cat_map.get(c, "zz"), c.lower()))
    
    candidates = sorted_coverage["Candidate"].tolist()
    overall_scores = {row["Candidate"]: row["Overall"] for _, row in sorted_coverage.iterrows()}
    
    # Build multi-index for rows including Overall as first row
    # Overall row will use ("📊 Overall", "Overall Score") as the index tuple
    disp = pd.DataFrame(
        index=pd.MultiIndex.from_tuples(
            [("📊 Overall", "Overall Score")] + [(cat_map.get(c,"uncategorised"), c) for c in ordered], 
            names=["Category","Criterion"]
        ), 
        columns=["Average"] + candidates
    )
    
    # Fill Overall row (no average for Overall, just candidate scores)
    for cand in candidates:
        overall = overall_scores[cand]
        icon = "✅" if overall >= hi else ("⚠️" if overall >= lo else "⛔")
        disp.loc[("📊 Overall", "Overall Score"), cand] = f"{overall:.2f} {icon}"  # Always show score (score first for proper sorting)
    disp.loc[("📊 Overall", "Overall Score"), "Average"] = ""  # Blank average for Overall
    
    # Fill in the matrix values with icons (and optional scores) for criteria
    for crit in ordered:
        # Calculate average score for this criterion across all displayed candidates
        scores_for_crit = [float(mat.loc[crit, cand]) for cand in candidates]
        avg_score = sum(scores_for_crit) / len(scores_for_crit) if scores_for_crit else 0.0
        
        # Format average column (score first for proper sorting)
        avg_icon = "✅" if avg_score >= hi else ("⚠️" if avg_score >= lo else "⛔")
        disp.loc[(cat_map.get(crit,"uncategorised"), crit), "Average"] = f"{avg_score:.2f} {avg_icon}" if show_scores else avg_icon
        
        # Format candidate columns (score first for proper sorting)
        for c in candidates:
            val = float(mat.loc[crit, c])
            icon = "✅" if val >= hi else ("⚠️" if val >= lo else "⛔")
            disp.loc[(cat_map.get(crit,"uncategorised"), crit), c] = f"{val:.2f} {icon}" if show_scores else icon
    
    # Display matrix by category with expandable sections
    if expand_categories:
        st.dataframe(disp, width='stretch')
    else:
        # Group by category for collapsible display
        categories = disp.index.get_level_values(0).unique()
        for category in categories:
            cat_data = disp.loc[category]
            with st.expander(f"📁 {category}", expanded=False):
                st.dataframe(cat_data, width='stretch')
    
    # Evidence Explorer - more prominent styling
    st.markdown("---")
    st.markdown("#### 🔍 Evidence Explorer")
    st.caption("Select a candidate and criterion to view the evidence snippet used for scoring")
    
    names = coverage_df["Candidate"].tolist()
    ordered = ordered  # ensure exists
    
    ev_col1, ev_col2 = st.columns(2)
    with ev_col1:
        cand_choice = st.selectbox("👤 Candidate", names, index=0 if names else None)
    with ev_col2:
        crit = st.selectbox("📋 Criterion", ordered, index=0 if ordered else None)
    
    if crit and cand_choice:
        ev_map = st.session_state.get("evidence_map", {})
        snip, _ = ev_map.get((cand_choice, crit), (None, None))
        
        # Get the actual score from coverage_df, not evidence_map
        actual_score = None
        if crit in coverage_df.columns:
            cand_row = coverage_df[coverage_df["Candidate"] == cand_choice]
            if not cand_row.empty:
                actual_score = float(cand_row[crit].iloc[0])
        
        if snip is None or actual_score is None:
            st.info("No cached evidence yet for this pair. Re-run Analyse to refresh evidence.")
        else:
            # Show score with color coding (using actual score from coverage_df)
            score_color = "#28a745" if actual_score >= hi else ("#ffc107" if actual_score >= lo else "#dc3545")
            st.markdown(f"**Score:** <span style='color: {score_color}; font-size: 1.2rem; font-weight: bold;'>{actual_score:.2f}</span>", unsafe_allow_html=True)
            st.markdown("**Evidence from resume:**")
            st.code(snip, language=None)

elif current_page == "Candidate Insights":
    st.markdown("<div style='margin-top: 1.5rem;'></div>", unsafe_allow_html=True)
    st.subheader("Candidate Insights")
    
    if not API_KEY_SET:
        st.error("OpenAI API key is required for Insights. Set `OPENAI_API_KEY` to enable this tab.")
        st.stop()
    
    insights_map = st.session_state.get("last_insights", {})
    covdf = st.session_state.get("last_coverage", pd.DataFrame())
    
    if not insights_map or (isinstance(covdf, pd.DataFrame) and covdf.empty):
        st.info("Run analysis first to view candidate insights.")
        st.stop()
    
    # Get threshold settings for color coding
    hi = st.session_state.get("cov_hi", 0.70)
    lo = st.session_state.get("cov_lo", 0.45)
    cat_map = st.session_state.get("cat_map", {})
    
    # Sort candidates by Overall score (descending) and add rank
    covdf_sorted = covdf.sort_values("Overall", ascending=False).reset_index(drop=True)
    covdf_sorted["Rank"] = range(1, len(covdf_sorted) + 1)
    
    # Build candidate options with rank and score: "Name (Score: 0.85) ⭐ #1"
    candidate_options = []
    for _, row in covdf_sorted.iterrows():
        name = row["Candidate"]
        score = row["Overall"]
        rank = row["Rank"]
        option_text = f"{name} (Score: {score:.2f}) ⭐ #{rank}"
        candidate_options.append(option_text)
    
    # Initialize or get the current selection index
    if "insights_selected_idx" not in st.session_state:
        st.session_state.insights_selected_idx = 0
    
    # Compact selector with navigation and download options
    selector_col1, nav_col1, nav_col2, download_col = st.columns([4, 1, 1, 1])
    
    with nav_col1:
        if st.button("⬅️ Previous", disabled=(st.session_state.insights_selected_idx == 0), use_container_width=True):
            st.session_state.insights_selected_idx = max(0, st.session_state.insights_selected_idx - 1)
            st.rerun()
    
    with nav_col2:
        if st.button("Next ➡️", disabled=(st.session_state.insights_selected_idx == len(candidate_options) - 1), use_container_width=True):
            st.session_state.insights_selected_idx = min(len(candidate_options) - 1, st.session_state.insights_selected_idx + 1)
            st.rerun()
    
    with download_col:
        # Initialize download trigger state if needed
        if "insights_download_trigger" not in st.session_state:
            st.session_state.insights_download_trigger = None
        
        download_format = st.selectbox(
            "Download",
            options=["📥 Report", "PDF", "Word"],
            key="insights_download_format",
            label_visibility="collapsed",
            help="Download report for current candidate"
        )
        
        # Trigger download when user selects PDF or Word
        if download_format not in ["📥 Report"] and download_format != st.session_state.insights_download_trigger:
            st.session_state.insights_download_trigger = download_format
    
    # Candidate selector with search (selectbox is searchable by default)
    with selector_col1:
        selected_option = st.selectbox(
            "Select a candidate to view their detailed insights",
            options=candidate_options,
            index=st.session_state.insights_selected_idx,
            help="Search by name or browse by rank. Sorted by overall score (best first)."
        )
    
    # Update the index if user manually selected from dropdown
    selected_idx = candidate_options.index(selected_option)
    if selected_idx != st.session_state.insights_selected_idx:
        st.session_state.insights_selected_idx = selected_idx
    
    # Use the session state index to get candidate info (this is the source of truth)
    current_idx = st.session_state.insights_selected_idx
    selected_name = covdf_sorted.iloc[current_idx]["Candidate"]
    selected_score = covdf_sorted.iloc[current_idx]["Overall"]
    selected_rank = covdf_sorted.iloc[current_idx]["Rank"]
    
    # Calculate criteria breakdown
    crit_cols = [c for c in covdf.columns if c not in ("Candidate", "Overall")]
    selected_row = covdf[covdf["Candidate"] == selected_name].iloc[0]
    strong_count = sum(1 for c in crit_cols if selected_row[c] >= hi)
    moderate_count = sum(1 for c in crit_cols if lo <= selected_row[c] < hi)
    weak_count = sum(1 for c in crit_cols if selected_row[c] < lo)
    
    # Get insights for this candidate
    info = insights_map.get(selected_name, {})
    
    # Get job title for reports
    jd_text = st.session_state.jd.text if st.session_state.get("jd") else ""
    job_title = ""
    if jd_text:
        jd_lines = [line.strip() for line in jd_text.split('\n') if line.strip()]
        for line in jd_lines[:5]:
            if len(line) < 10 or 'http' in line.lower() or any(char.isdigit() for char in line[:4]):
                continue
            if any(indicator in line.lower() for indicator in ['position:', 'role:', 'job title:', 'title:']):
                job_title = line.split(':', 1)[1].strip() if ':' in line else line
                break
            elif len(line) > 15 and len(line) < 100:
                job_title = line
                break
    
    # Generate download if format was just selected
    evidence_map = st.session_state.get("evidence_map", {})
    
    if st.session_state.insights_download_trigger == "PDF":
        if reportlab is not None:
            pdf_bytes = to_individual_candidate_pdf(
                selected_name, selected_row, info, evidence_map,
                cat_map, hi, lo, include_evidence=False, job_title=job_title
            )
            if pdf_bytes:
                with download_col:
                    st.download_button(
                        "Download",
                        data=pdf_bytes,
                        file_name=f"{selected_name.replace(' ', '_')}_report.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        key="insights_pdf_download"
                    )
                # Reset trigger after generating download button
                st.session_state.insights_download_trigger = None
        else:
            with download_col:
                st.warning("⚠️", icon="⚠️")
    
    elif st.session_state.insights_download_trigger == "Word":
        if docx is not None:
            docx_bytes = to_individual_candidate_docx(
                selected_name, selected_row, info, cat_map, hi, lo, job_title=job_title
            )
            if docx_bytes:
                with download_col:
                    st.download_button(
                        "Download",
                        data=docx_bytes,
                        file_name=f"{selected_name.replace(' ', '_')}_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="insights_docx_download"
                    )
                # Reset trigger after generating download button
                st.session_state.insights_download_trigger = None
        else:
            with download_col:
                st.warning("⚠️", icon="⚠️")
    
    # Compact header with all key info in one line
    score_color = "#28a745" if selected_score >= hi else ("#ffc107" if selected_score >= lo else "#dc3545")
    st.markdown(f"""
    <div style='background-color: rgba(128,128,128,0.1); padding: 1rem; border-radius: 0.5rem; margin: 1rem 0;'>
        <h3 style='margin: 0 0 0.5rem 0;'>👤 {selected_name} — At a Glance</h3>
        <div style='display: flex; gap: 2rem; align-items: center;'>
            <div>
                <span style='font-size: 0.9rem; color: #666;'>Score:</span>
                <span style='color: {score_color}; font-size: 1.8rem; font-weight: bold; margin-left: 0.5rem;'>{selected_score:.2f}</span>
            </div>
            <div>
                <span style='font-size: 0.9rem; color: #666;'>Rank:</span>
                <span style='font-size: 1.8rem; font-weight: bold; margin-left: 0.5rem;'>#{selected_rank}</span>
                <span style='font-size: 0.85rem; color: #666;'> of {len(covdf_sorted)}</span>
            </div>
            <div>
                <span style='font-size: 0.9rem; color: #666;'>Criteria:</span>
                <span style='margin-left: 0.5rem;'>✅ {strong_count}</span>
                <span style='margin-left: 0.5rem;'>⚠️ {moderate_count}</span>
                <span style='margin-left: 0.5rem;'>⛔ {weak_count}</span>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # ========== GPT INSIGHTS (MAIN FEATURE) ==========
    st.markdown("### 💡 GPT-Powered Insights")
    
    info = insights_map.get(selected_name, {})
    
    insight_col1, insight_col2 = st.columns([1, 1])
    
    with insight_col1:
        st.markdown("#### ✅ Top Strengths")
        strengths = info.get("top", [])
        if strengths:
            for strength in strengths:
                st.markdown(f"- {strength}")
        else:
            st.info("No strengths identified yet. Click 'Refresh GPT Insights' above.")
    
    with insight_col2:
        st.markdown("#### ⚠️ Gaps / Risks")
        gaps = info.get("gaps", [])
        if gaps:
            for gap in gaps:
                st.markdown(f"- {gap}")
        else:
            st.info("No gaps identified yet. Click 'Refresh GPT Insights' above.")
    
    # Additional notes
    notes = info.get("notes", "")
    if notes:
        st.markdown("#### 📝 Additional Notes")
        st.info(notes)
    
    # ========== DETAILED SCORES ==========
    with st.expander("📈 Detailed Scores for All Criteria", expanded=False):
        st.markdown(
            "View individual criteria scores with color-coding "
            "(<span style='color: #28a745; font-weight: bold;'>strong</span>  -  "
            "<span style='color: #ffc107; font-weight: bold;'>moderate</span>  -  "
            "<span style='color: #dc3545; font-weight: bold;'>weak or missing</span>)",
            unsafe_allow_html=True
        )
        
        # Build scores dataframe for this candidate
        scores_data = []
        for crit in crit_cols:
            score = selected_row[crit]
            category = cat_map.get(crit, "Uncategorized")
            scores_data.append({"Category": category, "Criterion": crit, "Score": score})
        
        scores_df = pd.DataFrame(scores_data)
        
        # Group by category if categories exist
        if cat_map:
            st.markdown("**Scores grouped by category:**")
            for category in sorted(scores_df["Category"].unique()):
                cat_scores = scores_df[scores_df["Category"] == category]
                st.markdown(f"**📁 {category}**")
                for _, row in cat_scores.iterrows():
                    score = row["Score"]
                    crit = row["Criterion"]
                    color = "#28a745" if score >= hi else ("#ffc107" if score >= lo else "#dc3545")
                    st.markdown(f"- {crit}: <span style='color: {color}; font-weight: bold;'>{score:.2f}</span>", unsafe_allow_html=True)
                st.markdown("")
        else:
            # No categories, just list all
            for _, row in scores_df.iterrows():
                score = row["Score"]
                crit = row["Criterion"]
                color = "#28a745" if score >= hi else ("#ffc107" if score >= lo else "#dc3545")
                st.markdown(f"- {crit}: <span style='color: {color}; font-weight: bold;'>{score:.2f}</span>", unsafe_allow_html=True)
        
        # Show evidence snippets inline
        st.markdown("---")
        st.markdown("**🔍 View Evidence for Specific Criterion**")
        
        ev_crit = st.selectbox("Select criterion to see evidence snippet", crit_cols, key="insights_evidence_crit")
        
        if ev_crit:
            ev_map = st.session_state.get("evidence_map", {})
            snip, _ = ev_map.get((selected_name, ev_crit), (None, None))
            actual_score = selected_row[ev_crit]
            
            if snip:
                score_color_ev = "#28a745" if actual_score >= hi else ("#ffc107" if actual_score >= lo else "#dc3545")
                st.markdown(f"**Score:** <span style='color: {score_color_ev}; font-size: 1.2rem; font-weight: bold;'>{actual_score:.2f}</span>", unsafe_allow_html=True)
                st.markdown("**Evidence from resume:**")
                st.code(snip, language=None)
            else:
                st.info("No cached evidence for this criterion. Re-run analysis to refresh.")
        
        # Collapse helper at bottom
        st.markdown("---")
        st.markdown("<div style='text-align: center; color: #666; font-size: 0.9rem;'>👆 Scroll to top to collapse this section</div>", unsafe_allow_html=True)
    
    # ========== RESUME ==========
    with st.expander("📄 Resume / CV", expanded=False):
        st.markdown("View the candidate's resume - both original file and extracted text")
        
        # Find the candidate object to get resume text
        cands = st.session_state.get("cached_candidates", [])
        cand_obj = None
        for c in cands:
            if c.name == selected_name:
                cand_obj = c
                break
        
        if cand_obj:
            st.markdown(f"**Filename:** {cand_obj.file_name}")
            
            # Show original file if it's a PDF
            if cand_obj.file_name.lower().endswith('.pdf'):
                with st.expander("📄 View Original PDF File", expanded=False):
                    import base64
                    base64_pdf = base64.b64encode(cand_obj.raw_bytes).decode('utf-8')
                    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600" type="application/pdf"></iframe>'
                    st.markdown(pdf_display, unsafe_allow_html=True)
                    st.caption("💡 Use the browser's PDF controls to zoom, navigate pages, or download if needed.")
            elif cand_obj.file_name.lower().endswith(('.docx', '.doc')):
                with st.expander("📥 Original DOCX/DOC File", expanded=False):
                    st.info("DOCX/DOC files cannot be previewed inline. View the 'Extracted Text' below, or download the file to open in Word.", icon="ℹ️")
                    # Provide download button for DOCX
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if cand_obj.file_name.lower().endswith('.docx') else "application/msword"
                    st.download_button(
                        "📥 Download to View in Word",
                        data=cand_obj.raw_bytes,
                        file_name=cand_obj.file_name,
                        mime=mime_type,
                        help="Download to view in Microsoft Word"
                    )
            else:
                # Plain text files in expander
                with st.expander("📄 View Original Text File", expanded=False):
                    st.code(cand_obj.text, language=None, line_numbers=False)
            
            st.markdown("---")
            st.markdown("**📝 Extracted Text**")
            st.code(cand_obj.text, language=None)
        else:
            st.warning("Resume text not available for this candidate.")
        
        # Collapse helper at bottom
        st.markdown("---")
        st.markdown("<div style='text-align: center; color: #666; font-size: 0.9rem;'>👆 Scroll to top to collapse this section</div>", unsafe_allow_html=True)

elif current_page == "Compare":
    st.markdown("<div style='margin-top: 1.5rem;'></div>", unsafe_allow_html=True)
    st.subheader("Compare Candidates vs JD")
    covdf = st.session_state.get("last_coverage", pd.DataFrame())
    names = covdf["Candidate"].tolist() if isinstance(covdf, pd.DataFrame) and not covdf.empty else []
    if not names:
        st.info("Run analysis to compare candidates.")
    else:
        sel = st.multiselect("Select 2–3 candidates", names, default=names[:2])
        if len(sel) < 2:
            st.warning("Select at least two candidates.")
        else:
            sub = covdf[covdf["Candidate"].isin(sel)].set_index("Candidate")
            st.write("**Comparison (scores)**")
            st.table(sub.reset_index().style.format({c: "{:.2f}" for c in sub.columns if c != "Candidate"}))

elif current_page == "Candidates":
    st.markdown("<div style='margin-top: 1.5rem;'></div>", unsafe_allow_html=True)
    st.subheader("Candidates")
    cands = st.session_state.get("cached_candidates", [])
    if not cands:
        st.info("Upload candidate resumes to view details.")
    else:
        # Build display options: "Label — Filename"
        options = [f"{c.name} — {c.file_name}" for c in cands]

        # Stronger, clearer header for the selector
        st.markdown(
            "<div class='cand-select-header'><strong>Select a resume</strong></div>",
            unsafe_allow_html=True
        )
        sel = st.selectbox("Select a resume", options=options, index=0, key="cand_selector", label_visibility="collapsed")
        sel_idx = options.index(sel) if sel in options else 0
        cand = cands[sel_idx]

        # CSS tweaks to improve look & spacing and align heights
        st.markdown(
            """
            <style>
            .cand-select-header { font-size: 1.05rem; font-weight: 700; margin: 0 0 0.25rem 0; }
            .readonly-input, .readonly-textarea {
                width: 100%;
                box-sizing: border-box;
                border: 1px solid rgba(49,51,63,0.2);
                border-radius: 0.5rem;
                padding: 0.5rem 0.75rem;
                background: var(--background-color);
                color: inherit;
                font: inherit;
            }
            /* Match Streamlit text_input height more closely */
            .readonly-input { height: 2.5rem; line-height: 1.2; }
            .readonly-textarea { height: 320px; resize: vertical; }
            .resume-section-gap { height: 0.75rem; }
            .field-label { font-weight: 700; margin-bottom: 0.25rem; display:block; }
            </style>
            """,
            unsafe_allow_html=True
        )

        col1, col2 = st.columns([1,1])
        with col1:
            st.markdown("<span class='field-label'>Filename</span>", unsafe_allow_html=True)
            st.markdown(
                f"<input class='readonly-input' type='text' value='{html.escape(cand.file_name)}' readonly/>",
                unsafe_allow_html=True
            )
        with col2:
            st.markdown("<span class='field-label'>Label (Edit if desired)</span>", unsafe_allow_html=True)
            lbl_key = f"cand_label_{cand.hash[:12]}"
            # Collapse the label on the Streamlit input (we show our own bold heading above)
            new_label = st.text_input("Label (Edit if desired)", value=cand.name, key=lbl_key, label_visibility="collapsed")
            if new_label != cand.name:
                cand.name = new_label

        # Add a little extra gap before the resume text
        st.markdown("<div class='resume-section-gap'></div>", unsafe_allow_html=True)
        st.markdown("<span class='field-label'>Resume Text Extracted</span>", unsafe_allow_html=True)
        st.code(cand.text)

        # Persist any label edits back to session_state
        st.session_state.cached_candidates = cands

elif current_page == "Scoring Explained":
    st.markdown("<div style='margin-top: 1.5rem;'></div>", unsafe_allow_html=True)
    st.subheader("📊 Scoring Explained")
    st.markdown("### What do the scores mean?")
    st.markdown("""
Each score is a **semantic similarity** between a JD **criterion** and the **most relevant snippet** in a candidate's resume.
It ranges from **0.00** (not similar) to **1.00** (nearly identical in meaning). It is **not** a percentage of criteria met.

**Color Coding:**
- ✅ **Green (Strong):** ≥0.70 — Excellent to good alignment
- ⚠️ **Yellow (Moderate):** 0.45–0.70 — Partial or indirect match
- ⛔ **Red (Weak):** <0.45 — Likely not covered

**Quick ranges:**
- **0.75–1.00:** Excellent alignment
- **0.60–0.74:** Good
- **0.45–0.59:** Moderate (partial/indirect)
- **0.30–0.44:** Weak
- **<0.30:** Likely not covered

**Overall** is a weighted average of all criteria scores.
""")

    st.markdown("### How it’s calculated (plain English)")
    st.markdown("""
1) The resume is split into chunks (e.g., ~1,200 chars).
2) For each **criterion**, we compare it to every chunk and keep the **best match**.
3) The similarity comes from vector math (cosine similarity on embeddings/TF‑IDF), which captures meaning overlap, not just the same words.
""")

    st.markdown("### Examples")
    st.markdown("""
- **Criterion:** “Project management with Agile/Scrum.”  
  **Resume:** “Led a 10‑person team delivering sprints using Scrum ceremonies.”  
  → **Score ≈ 0.75–0.85** (strong, direct alignment)

- **Criterion:** “Python data analysis.”  
  **Resume:** “Automated Excel reports; some VBA macros.”  
  → **Score ≈ 0.35–0.50** (partial/weak — tooling mismatch)

- **Criterion:** “Stakeholder communication.”  
  **Resume:** “Presented monthly performance updates to executives and coordinated cross‑team work.”  
  → **Score ≈ 0.60–0.75** (good evidence)
""")

    st.markdown("### How to use the scores")
    st.markdown("""
- Treat scores as **evidence strength**, not pass/fail. Use the Coverage matrix to spot strengths and gaps.
- Adjust **weights** to emphasise must‑haves; Overall will reflect your priorities.
- Use the **Evidence Explorer** (Coverage tab) to view the exact resume snippet behind each score.
""")

elif current_page == "Help & FAQ":
    st.markdown("<div style='margin-top: 1.5rem;'></div>", unsafe_allow_html=True)
    st.subheader("❓ Help & FAQ")
    
    # === SECTION 1: Scoring Explained ===
    with st.expander("📊 Scoring Explained", expanded=False):
        st.markdown("### What do the scores mean?")
        st.markdown("""
Each score is a **semantic similarity** between a JD **criterion** and the **most relevant snippet** in a candidate's resume.
It ranges from **0.00** (not similar) to **1.00** (nearly identical in meaning). It is **not** a percentage of criteria met.

**Color Coding:**
- ✅ **Green (Strong):** ≥0.70 — Excellent to good alignment
- ⚠️ **Yellow (Moderate):** 0.45–0.70 — Partial or indirect match
- ⛔ **Red (Weak):** <0.45 — Likely not covered

**Quick ranges:**
- **0.75–1.00:** Excellent alignment
- **0.60–0.74:** Good
- **0.45–0.59:** Moderate (partial/indirect)
- **0.30–0.44:** Weak
- **<0.30:** Likely not covered

**Overall** is a weighted average of all criteria scores.
""")

        st.markdown("### How it's calculated (plain English)")
        st.markdown("""
1) The resume is split into chunks (e.g., ~1,200 chars).
2) For each **criterion**, we compare it to every chunk and keep the **best match**.
3) The similarity comes from vector math (cosine similarity on embeddings/TF‑IDF), which captures meaning overlap, not just the same words.
""")

        st.markdown("### Examples")
        st.markdown("""
- **Criterion:** "Project management with Agile/Scrum."  
  **Resume:** "Led a 10‑person team delivering sprints using Scrum ceremonies."  
  → **Score ≈ 0.75–0.85** (strong, direct alignment)

- **Criterion:** "Python data analysis."  
  **Resume:** "Automated Excel reports; some VBA macros."  
  → **Score ≈ 0.35–0.50** (partial/weak — tooling mismatch)

- **Criterion:** "Stakeholder communication."  
  **Resume:** "Presented monthly performance updates to executives and coordinated cross‑team work."  
  → **Score ≈ 0.60–0.75** (good evidence)
""")

        st.markdown("### How to use the scores")
        st.markdown("""
- Treat scores as **evidence strength**, not pass/fail. Use the Coverage matrix to spot strengths and gaps.
- Review the **Scoring Analysis** page to see the full coverage matrix with color-coded results.
- Use the **Evidence Explorer** (Scoring Analysis tab) to view the exact resume snippet behind each score.
- Adjust scoring thresholds in **Settings** if you want stricter or more lenient classifications.
""")
    
    # === SECTION 2: Privacy & Data Handling ===
    with st.expander("🔒 Privacy & Data Handling", expanded=False):
        st.markdown("### How your data is handled")
        st.markdown("""
This application uses OpenAI's **GPT-4o** model for intelligent analysis of job descriptions and candidate resumes.

**Important privacy information:**

✅ **Your data is NOT used to train AI models**
- OpenAI does not use data submitted via their API to train or improve their models (unless you explicitly opt in)
- We have NOT opted in to any training programs

✅ **Limited data retention**
- API data is retained by OpenAI for up to **30 days** for abuse and misuse monitoring
- After 30 days, your data is automatically deleted from OpenAI's servers

✅ **Secure transmission**
- All data is transmitted securely via HTTPS to OpenAI's API endpoints
- No data is stored on external servers by this application

⚠️ **What this means for you:**
- Job descriptions and resumes are temporarily processed by OpenAI's systems
- If your organization has strict data privacy requirements, consult your compliance team before use
- If this data handling is unacceptable for your use case, do not use this application

📚 **Learn more:**
- [OpenAI API Data Usage Policies](https://openai.com/policies/api-data-usage-policies)
- [OpenAI Enterprise Privacy](https://openai.com/enterprise-privacy)
""")
    
    # === SECTION 3: PDF Text Extraction ===
    with st.expander("📄 PDF Text Extraction & OCR Tool", expanded=False):
        st.markdown("### PDF Extraction Settings")
        
        st.markdown("""If PDFs are extracting with text in the wrong order or appearing jumbled:""")
        
        # Layout-aware PDF toggle
        st.session_state["use_layout_pdf"] = st.checkbox(
            "Use layout-aware PDF extraction",
            value=st.session_state.get("use_layout_pdf", True),
            help="Disable if PDFs extract with jumbled order; falls back to standard extractor.",
            key="help_layout_pdf"
        )
        
        if st.session_state.get("use_layout_pdf", True):
            st.success("✅ Layout-aware extraction is enabled (recommended)")
        else:
            st.warning("⚠️ Using basic extraction mode. Re-upload files to apply.")
        
        st.markdown("""
**PDF extraction tips:**
- Ensure PDFs are text-based, not scanned images
- Very complex multi-column layouts may extract better with layout-aware mode disabled
- For scanned/image-based PDFs, use the OCR tool below
""")
        
        st.markdown("---")
        st.markdown("### 🔍 OCR Tool for Scanned PDFs")
        st.markdown("""
Use this tool when a PDF is image-based or poorly extracted. It will use **OCR (Optical Character Recognition)** 
to extract text from scanned documents.

**What you'll get:**
- 📝 **TXT file** - Plain text version of the extracted content
- 📄 **Simple PDF** - New text-based PDF (plain format, not layout-preserving)
""")
        
        # OCR Tool (moved from PDF Tools page)
        tesseract_ok = False
        tess_version = ""
        try:
            if pytesseract is not None:
                tess_version = str(pytesseract.get_tesseract_version())
                tesseract_ok = True
        except Exception:
            tesseract_ok = False
        
        if not tesseract_ok:
            st.warning("⚠️ Tesseract OCR not installed. This tool requires Tesseract to be installed on your system.")
            with st.expander("How to install Tesseract"):
                st.markdown("""
**Windows:**
1. Download from: https://github.com/UB-Mannheim/tesseract/wiki
2. Install and add to PATH
3. Restart this application

**Mac:** `brew install tesseract`

**Linux:** `sudo apt-get install tesseract-ocr`
""")
        else:
            with st.expander("ℹ️ System Information"):
                st.caption(f"**Tesseract OCR:** {tess_version}")
                st.caption(f"**Backends:** PyMuPDF {'✅' if fitz else '❌'} • pdf2image {'✅' if pdf2image else '❌'}")
            
            strong_ocr = st.checkbox("Use stronger OCR (for faint or blurry scans)", value=False,
                                    help="Renders at higher DPI (400 vs 300) with stronger contrast. Slower but better for poor quality scans.")
            
            conv_files = st.file_uploader("Upload PDFs to convert with OCR", type=["pdf"], accept_multiple_files=True, key="help_pdf_ocr_upl")
            
            if conv_files:
                for uf in conv_files:
                    st.markdown(f"**Processing: {uf.name}**")
                    b = uf.read()
                    
                    with st.spinner("Processing PDF..."):
                        try:
                            # Pass 1: Try direct text extraction
                            text_direct = ""
                            if fitz is not None:
                                with fitz.open(stream=b, filetype="pdf") as doc:
                                    chunks = []
                                    for p in doc:
                                        chunks.append(p.get_text("text"))
                                    text_direct = "\n".join(chunks)
                            
                            initial_length = len((text_direct or "").strip())
                            
                            # Pass 2: OCR if text is minimal
                            if initial_length < 200:
                                ocr_chunks = []
                                dpi = 400 if strong_ocr else 300
                                thresh_cut = 170 if strong_ocr else 200
                                tess_cfg = "--oem 3 --psm 6"
                                
                                used_backend = None
                                pages = None
                                
                                if pdf2image is not None:
                                    try:
                                        from pdf2image import convert_from_bytes
                                        pages = convert_from_bytes(b, dpi=dpi)
                                        used_backend = "pdf2image"
                                    except Exception as pdf2img_err:
                                        if "poppler" in str(pdf2img_err).lower():
                                            pages = None
                                        else:
                                            raise
                                
                                if pages is None and fitz is not None:
                                    pages = []
                                    with fitz.open(stream=b, filetype="pdf") as doc:
                                        mat = fitz.Matrix(dpi/72.0, dpi/72.0)
                                        for page in doc:
                                            pix = page.get_pixmap(matrix=mat, alpha=False)
                                            if Image is not None:
                                                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                                                pages.append(img)
                                    used_backend = "PyMuPDF"
                                
                                if pages:
                                    progress_bar = st.progress(0)
                                    status_text = st.empty()
                                    
                                    for idx, img in enumerate(pages):
                                        status_text.text(f"Processing page {idx+1} of {len(pages)}...")
                                        progress_bar.progress((idx + 1) / len(pages))
                                        
                                        if Image is not None:
                                            gray = img.convert("L")
                                            bw = gray.point(lambda x: 0 if x < thresh_cut else 255, "1").convert("L")
                                            img_for_ocr = bw
                                        else:
                                            img_for_ocr = img
                                        
                                        page_text = pytesseract.image_to_string(img_for_ocr, config=tess_cfg)
                                        ocr_chunks.append(page_text)
                                    
                                    progress_bar.empty()
                                    status_text.empty()
                                    text_ocr = "\n".join([t for t in ocr_chunks if t]).strip()
                                    
                                    if text_ocr:
                                        st.success(f"✅ Extracted {len(text_ocr)} characters from {len(pages)} page(s) using {used_backend}")
                                    else:
                                        st.warning("OCR completed but no text was extracted.")
                                        continue
                                else:
                                    st.error("Unable to rasterize PDF for OCR.")
                                    continue
                            else:
                                text_ocr = text_direct
                                st.info(f"ℹ️ PDF already contains {initial_length} characters of extractable text. OCR not needed.")
                            
                        except Exception as e:
                            st.error(f"Error during extraction: {str(e)}")
                            continue
                    
                    if not text_ocr:
                        st.error("No text could be extracted.")
                        continue
                    
                    # Download buttons
                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button(
                            "📥 Download TXT",
                            data=text_ocr.encode("utf-8", errors="ignore"),
                            file_name=uf.name.rsplit(".", 1)[0] + "_ocr.txt",
                            mime="text/plain",
                            use_container_width=True
                        )
                    with col2:
                        if reportlab is not None:
                            pdf_bytes = to_pdf_bytes_from_markdown(text_ocr)
                            if pdf_bytes:
                                st.download_button(
                                    "📥 Download PDF",
                                    data=pdf_bytes,
                                    file_name=uf.name.rsplit(".", 1)[0] + "_ocr.pdf",
                                    mime="application/pdf",
                                    use_container_width=True
                                )
                    
                    with st.expander("👁️ Preview extracted text"):
                        preview_text = text_ocr[:2000] + ("..." if len(text_ocr) > 2000 else "")
                        st.text(preview_text)
    
    # === SECTION 4: Troubleshooting ===
    with st.expander("🔧 Other Troubleshooting", expanded=False):
        st.markdown("### Common issues and solutions")
        
        st.markdown("#### Analysis not running")
        st.markdown("""
- Ensure you've uploaded both a Job Description and at least one candidate resume
- Check that criteria have been extracted (visible on Job Criteria page)
- Verify OpenAI API key is set if using GPT features
""")
        
        st.markdown("#### Scores seem too low/high")
        st.markdown("""
- Adjust scoring thresholds in **Settings** to match your standards
- Review criteria on **Job Criteria** page - overly specific criteria may score lower
- Check the **Evidence Explorer** to see what text is being matched
""")
    
    # === SECTION 5: Frequently Asked Questions ===
    with st.expander("💬 Frequently Asked Questions", expanded=False):
        st.markdown("### Can I use this without an OpenAI API key?")
        st.markdown("""
Partially. You can:
- Upload and view documents
- Manually create/edit criteria
- Run basic scoring analysis

You cannot:
- Auto-extract criteria from job descriptions
- Generate GPT-powered candidate insights
- Use intelligent recommendation features in reports
""")
        
        st.markdown("### How accurate is the scoring?")
        st.markdown("""
Scores measure **semantic similarity** based on text analysis. They are:
- Good at: Identifying relevant experience, skills, and qualifications mentioned in resumes
- Not perfect at: Understanding context, verifying claims, or assessing quality

**Best practice:** Use scores as a **screening tool** to identify strong candidates, then conduct interviews to verify fit.
""")
        
        st.markdown("### Can I edit candidate names or criteria?")
        st.markdown("""
**Yes!**
- Edit candidate names in the **Candidate Insights** page (name field is editable)
- Edit criteria on the **Job Criteria** page (inline editing or manual text entry)
- Changes persist for the current session
""")
        
        st.markdown("### What file formats are supported?")
        st.markdown("""
**Job Descriptions:** PDF, DOCX, DOC, TXT  
**Resumes:** PDF, DOCX, DOC, TXT  
**Cover Letters:** PDF, DOCX, DOC, TXT

All files are converted to text for analysis.
""")
        
        st.markdown("### How do I export results?")
        st.markdown("""
Visit the **Export Reports** page to download:
- Executive Summary (PDF or Word)
- Coverage Matrix (Excel with color coding, or CSV)
- Individual Candidate Reports (PDF or Word)
- Quick data exports (CSV)
""")
    
    # === SECTION 6: About ===
    with st.expander("ℹ️ About This Application", expanded=False):
        st.markdown("### Candidate Summariser")
        st.markdown("""
**Version:** 1.0 (November 2025)

**Technology:**
- **AI Model:** OpenAI GPT-4o
- **Embeddings:** Sentence transformers for semantic similarity
- **Framework:** Streamlit (Python)

**Key Features:**
- Automated criteria extraction from job descriptions
- Semantic similarity scoring between JD criteria and candidate resumes
- GPT-powered candidate insights and recommendations
- Professional report generation (PDF, Word, Excel)
- Cover letter integration and analysis

**Purpose:**
Streamline recruitment by automatically analyzing candidate fit against job requirements, 
helping hiring managers quickly identify top candidates and make data-driven decisions.
""")

elif current_page == "Export Reports":
    st.markdown("<div style='margin-top: 1.5rem;'></div>", unsafe_allow_html=True)
    st.subheader("Export Reports")
    
    covdf = st.session_state.get("last_coverage", pd.DataFrame())
    
    if not isinstance(covdf, pd.DataFrame) or covdf.empty:
        st.info("📊 Run analysis first to enable exports and reporting features.")
        st.stop()
    
    # Get necessary data
    insights = st.session_state.get("last_insights", {})
    jd_text = st.session_state.jd.text if st.session_state.get("jd") else ""
    cat_map = st.session_state.get("cat_map", {})
    evidence_map = st.session_state.get("evidence_map", {})
    hi = st.session_state.get("cov_hi", 0.70)
    lo = st.session_state.get("cov_lo", 0.45)
    
    # Overview metrics
    st.markdown("### 📊 Analysis Summary")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Candidates Analyzed", len(covdf))
    with col2:
        st.metric("Evaluation Criteria", len([c for c in covdf.columns if c not in ('Candidate', 'Overall')]))
    with col3:
        top_score = covdf['Overall'].max()
        st.metric("Highest Score", f"{top_score:.2f}")
    
    st.markdown("---")
    
    # === SECTION 1: Executive Summary ===
    st.markdown("### 📋 Executive Summary")
    st.caption("Concise overview for decision-makers and stakeholders")
    
    col_exec1, col_exec2 = st.columns([3, 1])
    
    with col_exec1:
        st.markdown("""
        **Includes:**
        - Analysis overview & key metrics
        - Top 5 candidates with scores
        - Key insights for top 3 candidates
        - Shortlist recommendation
        """)
    
    with col_exec2:
        if reportlab is not None:
            jd_filename = st.session_state.jd.file_name if st.session_state.get("jd") else "Job Description"
            exec_pdf = to_executive_summary_pdf(covdf, insights, jd_text, cat_map, hi, lo, jd_filename)
            if exec_pdf:
                st.download_button(
                    "📄 Download PDF",
                    data=exec_pdf,
                    file_name=f"executive_summary_{pd.Timestamp.now().strftime('%Y%m%d')}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    type="primary"
                )
            else:
                st.error("Failed to generate PDF")
        else:
            st.warning("Install 'reportlab' to enable PDF export")
        
        # Word export
        if docx is not None:
            jd_filename = st.session_state.jd.file_name if st.session_state.get("jd") else "Job Description"
            exec_docx = to_executive_summary_docx(covdf, insights, jd_text, cat_map, hi, lo, jd_filename)
            if exec_docx:
                st.download_button(
                    "📝 Download Word",
                    data=exec_docx,
                    file_name=f"executive_summary_{pd.Timestamp.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    help="Editable Word document"
                )
            else:
                st.error("Failed to generate Word document")
        else:
            st.caption("💡 python-docx available for Word export")
    
    st.markdown("---")
    
    # === SECTION 2: Detailed Analysis ===
    st.markdown("### 📊 Detailed Analysis Reports")
    st.caption("Comprehensive data for in-depth review")
    
    col_det1, col_det2 = st.columns([3, 1])
    
    with col_det1:
        st.markdown("""
        **Coverage Matrix Export:**
        - Full scores for all candidates × all criteria
        - Color-coded performance indicators
        - Summary statistics
        """)
    
    with col_det2:
        # Excel export
        if openpyxl is not None:
            excel_data = to_excel_coverage_matrix(covdf, cat_map, hi, lo)
            if excel_data:
                st.download_button(
                    "📊 Download Excel",
                    data=excel_data,
                    file_name=f"coverage_matrix_{pd.Timestamp.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    type="primary"
                )
            else:
                st.error("Failed to generate Excel")
        else:
            st.warning("Install 'openpyxl' for Excel export")
        
        # CSV fallback
        csv_data = covdf.to_csv(index=False).encode('utf-8')
        st.download_button(
            "📥 Download CSV",
            data=csv_data,
            file_name=f"coverage_matrix_{pd.Timestamp.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
            use_container_width=True
        )
    
    st.markdown("---")
    
    # === SECTION 3: Individual Candidate Reports ===
    st.markdown("### 👤 Individual Candidate Reports")
    st.caption("Detailed profile for each candidate - ideal for interview preparation")
    
    # Extract job title for reports
    job_title = ""
    if jd_text:
        jd_lines = [line.strip() for line in jd_text.split('\n') if line.strip()]
        for line in jd_lines[:5]:
            if len(line) < 10 or 'http' in line.lower() or any(char.isdigit() for char in line[:4]):
                continue
            if any(indicator in line.lower() for indicator in ['position:', 'role:', 'job title:', 'title:']):
                job_title = line.split(':', 1)[1].strip() if ':' in line else line
                break
            elif len(line) > 15 and len(line) < 100:
                job_title = line
                break
    
    # Selection options
    report_option = st.radio(
        "Select candidates for individual reports:",
        ["Top 3 Candidates", "Top 5 Candidates", "All Candidates", "Custom Selection"],
        horizontal=True
    )
    
    selected_candidates = []
    
    if report_option == "Top 3 Candidates":
        selected_candidates = covdf.head(3)['Candidate'].tolist()
    elif report_option == "Top 5 Candidates":
        selected_candidates = covdf.head(5)['Candidate'].tolist()
    elif report_option == "All Candidates":
        selected_candidates = covdf['Candidate'].tolist()
    else:  # Custom Selection
        all_candidates = covdf['Candidate'].tolist()
        selected_candidates = st.multiselect(
            "Choose specific candidates:",
            all_candidates,
            default=covdf.head(3)['Candidate'].tolist()
        )
    
    # Options
    include_evidence = st.checkbox(
        "Include evidence snippets (top 5 criteria per candidate)",
        value=False,
        help="Makes reports longer but provides context for scores"
    )
    
    if selected_candidates:
        st.info(f"📝 Ready to generate reports for **{len(selected_candidates)}** candidate(s)")
        
        col_ind1, col_ind2, col_ind3 = st.columns([2, 1, 1])
        
        with col_ind1:
            st.markdown(f"**Selected:** {', '.join(selected_candidates[:3])}" + 
                       (f" and {len(selected_candidates)-3} more..." if len(selected_candidates) > 3 else ""))
        
        with col_ind2:
            # Generate PDF reports
            if st.button("📄 PDF Reports", use_container_width=True, type="primary"):
                if reportlab is not None:
                    with st.spinner("Generating PDFs..."):
                        # Create a merged PDF with all candidates
                        from io import BytesIO
                        try:
                            from PyPDF2 import PdfMerger
                            merger = PdfMerger()
                            
                            for cand_name in selected_candidates:
                                cand_row = covdf[covdf['Candidate'] == cand_name].iloc[0]
                                cand_insights = insights.get(cand_name, {})
                                
                                pdf_bytes = to_individual_candidate_pdf(
                                    cand_name, cand_row, cand_insights, evidence_map, 
                                    cat_map, hi, lo, include_evidence, job_title=job_title
                                )
                                
                                if pdf_bytes:
                                    merger.append(BytesIO(pdf_bytes))
                            
                            output = BytesIO()
                            merger.write(output)
                            merger.close()
                            output.seek(0)
                            
                            st.download_button(
                                "⬇️ Download Combined PDF",
                                data=output.getvalue(),
                                file_name=f"candidate_reports_{pd.Timestamp.now().strftime('%Y%m%d')}.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                        except ImportError:
                            # Fallback: just generate first candidate if PyPDF2 not available
                            st.warning("Install 'PyPDF2' for combined PDFs. Showing first candidate only.")
                            cand_name = selected_candidates[0]
                            cand_row = covdf[covdf['Candidate'] == cand_name].iloc[0]
                            cand_insights = insights.get(cand_name, {})
                            
                            pdf_bytes = to_individual_candidate_pdf(
                                cand_name, cand_row, cand_insights, evidence_map,
                                cat_map, hi, lo, include_evidence, job_title=job_title
                            )
                            
                            if pdf_bytes:
                                st.download_button(
                                    "⬇️ Download PDF",
                                    data=pdf_bytes,
                                    file_name=f"{cand_name.replace(' ', '_')}_report.pdf",
                                    mime="application/pdf",
                                    use_container_width=True
                                )
                else:
                    st.warning("Install 'reportlab' to generate PDFs")
        
        with col_ind3:
            # Generate Word reports (editable)
            if st.button("📝 Word Reports", use_container_width=True):
                if docx is not None:
                    st.info("💡 Download individual Word documents below (fully editable):")
                    
                    for cand_name in selected_candidates:
                        cand_row = covdf[covdf['Candidate'] == cand_name].iloc[0]
                        cand_insights = insights.get(cand_name, {})
                        
                        docx_bytes = to_individual_candidate_docx(
                            cand_name, cand_row, cand_insights, cat_map, hi, lo, job_title=job_title
                        )
                        
                        if docx_bytes:
                            st.download_button(
                                f"📝 {cand_name}",
                                data=docx_bytes,
                                file_name=f"{cand_name.replace(' ', '_')}_report.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"docx_{cand_name}",
                                help="Editable Word document"
                            )
                else:
                    st.warning("Install 'python-docx' to generate Word documents")
    else:
        st.warning("⚠️ Select at least one candidate to generate reports")
    
    st.markdown("---")
    
    # === SECTION 4: Quick Data Exports ===
    with st.expander("📤 Quick Data Exports", expanded=False):
        st.markdown("**Export raw data for further analysis:**")
        
        quick_col1, quick_col2 = st.columns(2)
        
        with quick_col1:
            # Candidate list with scores
            candidate_summary = covdf[['Candidate', 'Overall']].copy()
            candidate_summary['Rank'] = range(1, len(candidate_summary) + 1)
            candidate_summary = candidate_summary[['Rank', 'Candidate', 'Overall']]
            
            csv_candidates = candidate_summary.to_csv(index=False).encode('utf-8')
            st.download_button(
                "📋 Candidate List (CSV)",
                data=csv_candidates,
                file_name="candidate_summary.csv",
                mime="text/csv",
                help="Simple list with rank and overall scores"
            )
        
        with quick_col2:
            # Criteria list
            criteria_list = [c for c in covdf.columns if c not in ('Candidate', 'Overall')]
            criteria_df = pd.DataFrame({
                'Criterion': criteria_list,
                'Category': [cat_map.get(c, 'Uncategorized') for c in criteria_list]
            })
            
            csv_criteria = criteria_df.to_csv(index=False).encode('utf-8')
            st.download_button(
                "📝 Criteria List (CSV)",
                data=csv_criteria,
                file_name="criteria_list.csv",
                mime="text/csv",
                help="List of all evaluation criteria with categories"
            )

# -----------------------------
# Settings Page
# -----------------------------
elif current_page == "Settings":
    st.markdown("<div style='margin-top: 1.5rem;'></div>", unsafe_allow_html=True)
    st.subheader("⚙️ Settings")

    # Scoring Thresholds Section
    st.markdown("### 📊 Scoring Thresholds")
    st.caption("These settings control how scores are displayed throughout the app (Coverage, Compare, etc.)")
    
    st.markdown("**Threshold Values:**")
    thresh_col1, thresh_col2 = st.columns(2)
    
    with thresh_col1:
        hi = st.slider(
            "✅ High threshold (green)", 
            0.5, 0.95, 
            st.session_state.get("cov_hi", 0.70), 
            step=0.01,
            help="Scores at or above this are marked with ✅ (green). Recommended: 0.65-0.75"
        )
        st.session_state.cov_hi = hi
    
    with thresh_col2:
        lo = st.slider(
            "⚠️ Medium threshold (yellow)", 
            0.1, float(hi-0.01), 
            st.session_state.get("cov_lo", 0.45), 
            step=0.01,
            help="Scores at or above this are marked with ⚠️ (yellow). Below this get ⛔ (red). Recommended: 0.40-0.50"
        )
        st.session_state.cov_lo = lo
    
    # Display current thresholds and reset button
    info_col, reset_col = st.columns([3, 1])
    with info_col:
        st.info(f"Current thresholds: ✅ High ≥ {hi:.2f} | ⚠️ Medium ≥ {lo:.2f} | ⛔ Low < {lo:.2f}")
    with reset_col:
        if st.button("Reset to Defaults", help="Reset thresholds to recommended defaults (High: 0.70, Medium: 0.45)"):
            st.session_state.cov_hi = 0.70
            st.session_state.cov_lo = 0.45
            st.rerun()
    
    st.markdown("---")
    st.markdown("### 💡 Need Help?")
    st.markdown("""
Visit the **Help & FAQ** page for:
- Detailed explanation of how scoring works
- Privacy and data handling information
- Troubleshooting tips (including PDF extraction issues)
- Frequently asked questions
""")

# -----------------------------
# Orchestration
# -----------------------------
if st.session_state.get("_trigger_analyse", False):
    # Get status container if it exists
    status_container = st.session_state.get("_status_container", st.empty())
    
    crit_text = st.session_state.get("criteria_text", "").strip()
    if not crit_text and st.session_state.get("jd") is not None:
        # default build if nothing present: GPT-only
        jd_txt = st.session_state.jd.text
        if _get_openai_client()[0] is None:
            st.error("GPT extraction requires an OpenAI API key. Set OPENAI_API_KEY or paste criteria manually on the JD tab.")
        else:
            secs = cached_extract_jd_sections_with_gpt(st.session_state.jd.hash, jd_txt)
            crits, cat_map = build_criteria_from_gpt_sections(secs, per_section=999, cap_total=10000)
            st.session_state.criteria_text = "\n".join(crits)
            st.session_state.cat_map = cat_map
            crit_text = st.session_state.criteria_text

    

    # Warn if JD text is suspiciously short — extraction may have failed
    try:
        jd_len = len(st.session_state.jd.text or "")
        if jd_len < 200:
            st.warning("JD text looks very short. Try enabling the robust PDF extractor or use 'Extract with GPT' to confirm sections.", icon="⚠️")
    except Exception:
        pass

    criteria = parse_criteria_text(crit_text)
    weights = get_weights(criteria, st.session_state.get("weights_mode","Uniform"), st.session_state.get("weights_csv",""))

    # --- Auto-build criteria if still empty (fallback on Analyze) ---
    if (not criteria) and st.session_state.get("jd"):
        jd_txt_fb = st.session_state.jd.text
        if _get_openai_client()[0] is not None:
            secs_fb = cached_extract_jd_sections_with_gpt(st.session_state.jd.hash, jd_txt_fb)
            crits_fb, cat_map_fb = build_criteria_from_gpt_sections(secs_fb, per_section=999, cap_total=10000)
        else:
            st.error("GPT extraction requires an OpenAI API key. Set OPENAI_API_KEY or paste criteria manually on the JD tab.")
            crits_fb, cat_map_fb = [], {}
        if crits_fb:
            st.session_state.criteria_text = "\n".join(crits_fb)
            st.session_state.cat_map = cat_map_fb
            criteria = parse_criteria_text(st.session_state.criteria_text)
            try:
                st.toast(f"Built {len(criteria)} criteria automatically.", icon="🧩")
            except Exception:
                pass

    if not criteria:
        st.warning("No criteria to analyse. Build criteria (Job Description tab) then click Analyse.")
    elif not st.session_state.get("cached_candidates"):
        st.warning("Please upload candidate resumes before running Analyze.")
    else:
        status_container.info("🔍 Extracting JD & building criteria...")
        
        with st.spinner("Running analysis…"):
            cov, ins_local, snips, ev_map = analyse_candidates(
                st.session_state.cached_candidates, criteria, weights,
                chunk_chars=st.session_state.get("chunk_chars",1200),
                overlap=st.session_state.get("overlap",150)
            )
            
            status_container.info("📊 Scoring candidates complete. Preparing insights...")
            
            st.session_state.last_coverage = cov
            st.session_state.last_insights = ins_local
            st.session_state.last_snippets = snips
            st.session_state.evidence_map = ev_map

                        # ---- GPT insight generation ----
            api_ok = _get_openai_client()[0] is not None

            if api_ok and isinstance(cov, pd.DataFrame) and not cov.empty:
                status_container.info("✨ Generating GPT insights...")
                try:
                    with st.spinner("Generating GPT insights…"):
                        jd_text_full = st.session_state.jd.text if st.session_state.get("jd") else ""
                        upgraded = {}
                        crit_cols = [c for c in cov.columns if c not in ("Candidate","Overall")]
                        for _, row in cov.iterrows():
                            cand_name = row["Candidate"]
                            # Row dict: {criterion: score}
                            row_dict = {c: float(row[c]) for c in crit_cols}
                            # Retrieve cached candidate text
                            cand_text = ""
                            for c_obj in st.session_state.get("cached_candidates", []):
                                if c_obj.name == cand_name:
                                    cand_text = c_obj.text
                                    break
                            upgraded[cand_name] = gpt_candidate_insights(
                                candidate_name=cand_name,
                                candidate_text=cand_text,
                                jd_text=jd_text_full,
                                criteria=crit_cols,
                                coverage_row=row_dict,
                                evidence_map=st.session_state.get("evidence_map", {}),
                            ) or st.session_state["last_insights"].get(cand_name, {})
                        # Replace insights with GPT versions where available
                        st.session_state["last_insights"] = upgraded
                        status_container.success("✅ Analysis complete!")
                        try: st.toast("Insights enhanced with GPT.", icon="✨")
                        except Exception: pass
                except Exception:
                    st.warning("GPT insight generation failed — using local insights.")
                    status_container.success("✅ Analysis complete (without GPT insights)")
            elif not api_ok:
                st.info("OpenAI API key not configured. Keeping baseline (non-GPT) insights.")
                status_container.success("✅ Analysis complete!")
            else:
                status_container.success("✅ Analysis complete!")
            
            # ---- Log usage for cost tracking ----
            try:
                jd_text = st.session_state.jd.text if st.session_state.get("jd") else ""
                candidate_texts = [c.text for c in st.session_state.cached_candidates]
                total_resume_chars = sum(len(t) for t in candidate_texts)
                
                cost_info = estimate_analysis_cost(jd_text, candidate_texts, len(criteria))
                
                log_usage(
                    num_candidates=len(st.session_state.cached_candidates),
                    num_criteria=len(criteria),
                    total_resume_chars=total_resume_chars,
                    jd_chars=len(jd_text),
                    estimated_cost=cost_info["estimated_cost_usd"],
                    model=cost_info["model"]
                )
            except Exception as e:
                # Don't fail analysis if logging fails
                pass
                
            st.session_state._trigger_analyse = False
            st.rerun()

        if isinstance(st.session_state.last_coverage, pd.DataFrame) and not st.session_state.last_coverage.empty:
            try: st.toast("Analysis complete.", icon="✅")
            except Exception: pass
        else:
            st.warning("Analysis ran but produced no rows. Check criteria and document text.")

# -----------------------------
# Insights-only refresh orchestration (no rescoring)
# -----------------------------
if st.session_state.get("_trigger_refresh_insights", False):
    cov = st.session_state.get("last_coverage", pd.DataFrame())
    if not isinstance(cov, pd.DataFrame) or cov.empty:
        st.warning("No coverage table available. Run full analysis first.")
        st.session_state._trigger_refresh_insights = False
    else:
        api_ok = _get_openai_client()[0] is not None
        if not api_ok:
            st.info("OpenAI API not configured. Set OPENAI_API_KEY to refresh GPT insights.")
            st.session_state._trigger_refresh_insights = False
        else:
            _update_stage("✨ Regenerating GPT insights…")
            try:
                with st.spinner("Regenerating GPT insights…"):
                    jd_text_full = st.session_state.jd.text if st.session_state.get("jd") else ""
                    upgraded = {}
                    crit_cols = [c for c in cov.columns if c not in ("Candidate","Overall")]
                    for _, row in cov.iterrows():
                        cand_name = row["Candidate"]
                        row_dict = {c: float(row[c]) for c in crit_cols}
                        cand_text = ""
                        for c_obj in st.session_state.get("cached_candidates", []):
                            if c_obj.name == cand_name:
                                cand_text = c_obj.text
                                break
                        upgraded[cand_name] = gpt_candidate_insights(
                            f"{cand_name}",
                            cand_text,
                            jd_text_full,
                            crit_cols,
                            row_dict,
                            st.session_state.get("evidence_map", {}),
                        ) or st.session_state.get("last_insights", {}).get(cand_name, {})
                    st.session_state["last_insights"] = upgraded
                    st.success("✅ GPT insights refreshed.")
            finally:
                st.session_state._trigger_refresh_insights = False
            st.rerun()