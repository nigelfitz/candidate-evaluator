"""
Export utilities for generating PDF, Excel, and Word reports
Ported from Streamlit app with full feature parity
"""

import io
import json
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
import pandas as pd

# Optional imports for export functionality
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.utils.dataframe import dataframe_to_rows
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False


# ============================================================================
# PDF PREVIEW UTILITIES
# ============================================================================

def render_pdf_to_images(pdf_bytes: bytes, max_pages: int = 10, dpi: int = 150) -> List[bytes]:
    """
    Render PDF pages as PNG images for web preview.
    Returns list of PNG image bytes (one per page).
    
    Args:
        pdf_bytes: PDF file as bytes
        max_pages: Maximum number of pages to render
        dpi: DPI for rendering (150 is good balance of quality/size)
    
    Returns:
        List of PNG image bytes, or empty list if PyMuPDF not available
    """
    if not PYMUPDF_AVAILABLE:
        return []
    
    try:
        images = []
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            total_pages = len(doc)
            pages_to_render = min(total_pages, max_pages)
            
            for page_num in range(pages_to_render):
                page = doc[page_num]
                # Render at specified DPI (150 DPI = 150/72 scale factor)
                pix = page.get_pixmap(matrix=fitz.Matrix(dpi/72, dpi/72))
                img_bytes = pix.tobytes("png")
                images.append(img_bytes)
        
        return images
    except Exception as e:
        print(f"Error rendering PDF to images: {e}")
        return []


# ============================================================================
# EXCEL EXPORT - Coverage Matrix with Color Coding
# ============================================================================

def to_excel_coverage_matrix(
    coverage: pd.DataFrame, 
    cat_map: Dict[str, str], 
    hi: float = 0.75, 
    lo: float = 0.35
) -> Optional[bytes]:
    """
    Export coverage matrix to Excel with formatting and color coding.
    
    Creates a multi-sheet workbook:
    - Sheet 1: Summary (metrics + top 5 candidates)
    - Sheet 2: Coverage Matrix (transposed: criteria as rows, candidates as columns)
    - Sheet 3: Legend (score color coding explanation)
    
    Args:
        coverage: DataFrame with columns ['Candidate', 'Overall', ...criteria...]
        cat_map: Dict mapping criterion -> category
        hi: High threshold for color coding (green)
        lo: Low threshold for color coding (red)
    
    Returns:
        Bytes of Excel file or None if openpyxl not available
    """
    if not OPENPYXL_AVAILABLE:
        return None
    
    try:
        buf = io.BytesIO()
        
        # Create workbook with multiple sheets
        wb = openpyxl.Workbook()
        
        # Sheet 1: Summary
        ws_summary = wb.active
        ws_summary.title = "Summary"
        
        # Add title
        ws_summary['A1'] = "Candidate Analysis Summary"
        ws_summary['A1'].font = Font(size=16, bold=True, color="1F77B4")
        ws_summary.merge_cells('A1:D1')
        
        # Add metrics
        row = 3
        ws_summary[f'A{row}'] = "Total Candidates:"
        ws_summary[f'B{row}'] = len(coverage)
        ws_summary[f'A{row}'].font = Font(bold=True)
        
        row += 1
        crit_cols = [c for c in coverage.columns if c not in ('Candidate', 'Overall')]
        ws_summary[f'A{row}'] = "Total Criteria:"
        ws_summary[f'B{row}'] = len(crit_cols)
        ws_summary[f'A{row}'].font = Font(bold=True)
        
        row += 2
        ws_summary[f'A{row}'] = "Top 5 Candidates"
        ws_summary[f'A{row}'].font = Font(size=14, bold=True)
        row += 1
        
        # Top candidates table
        top5 = coverage[['Candidate', 'Overall']].head(5)
        ws_summary[f'A{row}'] = "Rank"
        ws_summary[f'B{row}'] = "Candidate"
        ws_summary[f'C{row}'] = "Overall Score"
        for col in ['A', 'B', 'C']:
            ws_summary[f'{col}{row}'].font = Font(bold=True)
            ws_summary[f'{col}{row}'].fill = PatternFill(start_color="E7F3FF", end_color="E7F3FF", fill_type="solid")
        
        for idx, (_, r) in enumerate(top5.iterrows(), 1):
            row += 1
            ws_summary[f'A{row}'] = idx
            ws_summary[f'B{row}'] = r['Candidate']
            ws_summary[f'C{row}'] = round(r['Overall'], 2)
            
            # Color code the score
            score = r['Overall']
            if score >= hi:
                ws_summary[f'C{row}'].fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
            elif score >= lo:
                ws_summary[f'C{row}'].fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
            else:
                ws_summary[f'C{row}'].fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
        
        # Adjust column widths
        ws_summary.column_dimensions['A'].width = 8
        ws_summary.column_dimensions['B'].width = 30
        ws_summary.column_dimensions['C'].width = 15
        
        # Sheet 2: Coverage Matrix (TRANSPOSED)
        ws_matrix = wb.create_sheet("Coverage Matrix")
        
        # Get criteria columns (exclude Candidate and Overall)
        criteria_cols = [c for c in coverage.columns if c not in ('Candidate', 'Overall')]
        
        # Create transposed structure: Criteria as rows, Candidates as columns
        # Header row: "Criterion" then candidate names
        header_row = ['Criterion'] + coverage['Candidate'].tolist()
        
        # Write header
        for c_idx, value in enumerate(header_row, 1):
            cell = ws_matrix.cell(row=1, column=c_idx, value=value)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="1F77B4", end_color="1F77B4", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")
        
        # Write data rows (one criterion per row)
        for r_idx, criterion in enumerate(criteria_cols, 2):
            # First column: criterion name
            cell = ws_matrix.cell(row=r_idx, column=1, value=criterion)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="E7F3FF", end_color="E7F3FF", fill_type="solid")
            cell.alignment = Alignment(horizontal="left")
            
            # Remaining columns: scores for each candidate
            for c_idx, (_, candidate_row) in enumerate(coverage.iterrows(), 2):
                score = candidate_row[criterion]
                cell = ws_matrix.cell(row=r_idx, column=c_idx, value=round(score, 2))
                
                # Color code the score
                if score >= hi:
                    cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
                elif score >= lo:
                    cell.fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
                else:
                    cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
                cell.alignment = Alignment(horizontal="center")
        
        # Auto-adjust column widths
        for column in ws_matrix.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws_matrix.column_dimensions[column_letter].width = adjusted_width
        
        # Sheet 3: Legend
        ws_legend = wb.create_sheet("Legend")
        ws_legend['A1'] = "Score Color Coding"
        ws_legend['A1'].font = Font(size=14, bold=True)
        
        ws_legend['A3'] = f"Strong (≥{hi:.2f})"
        ws_legend['A3'].fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
        
        ws_legend['A4'] = f"Moderate ({lo:.2f}-{hi:.2f})"
        ws_legend['A4'].fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
        
        ws_legend['A5'] = f"Weak (<{lo:.2f})"
        ws_legend['A5'].fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
        
        ws_legend.column_dimensions['A'].width = 25
        
        wb.save(buf)
        buf.seek(0)
        return buf.getvalue()
        
    except Exception as e:
        print(f"Excel export error: {e}")
        return None


# ============================================================================
# PDF EXPORT - Executive Summary
# ============================================================================

def to_executive_summary_pdf(
    coverage: pd.DataFrame,
    insights: Dict[str, Dict],
    jd_text: str,
    cat_map: Dict[str, str],
    hi: float = 0.75,
    lo: float = 0.35,
    jd_filename: str = "Job Description",
    job_number: Optional[int] = None
) -> Optional[bytes]:
    """
    Generate a professional executive summary PDF.
    
    Includes:
    - Analysis overview & key metrics
    - Top 5 candidates with scores
    - AI insights for candidates (if available)
    - Shortlist recommendation
    
    Args:
        coverage: DataFrame with candidate scores
        insights: Dict mapping candidate_name -> {top, gaps, notes}
        jd_text: Job description text
        cat_map: Criterion -> category mapping
        hi: High threshold
        lo: Low threshold
        jd_filename: Name of JD file for display
        job_number: Analysis/Job ID number for header
    
    Returns:
        Bytes of PDF file or None if reportlab not available
    """
    if not REPORTLAB_AVAILABLE:
        return None
    
    try:
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#1F77B4'),
            spaceAfter=20,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#1F77B4'),
            spaceAfter=10,
            spaceBefore=10
        )
        
        # Title with Job Number
        if job_number:
            title_text = f"Executive Summary - Candidate Analysis<br/><font size=14>Job #{job_number:04d}: {jd_filename}</font>"
        else:
            title_text = f"Executive Summary - Candidate Analysis<br/><font size=14>Job Position: {jd_filename}</font>"
        
        story.append(Paragraph(title_text, title_style))
        
        # Generated date - subtle
        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#6B7280'),
            alignment=TA_CENTER
        )
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", date_style))
        story.append(Spacer(1, 0.5*cm))
        
        # Analysis Overview
        story.append(Paragraph("Analysis Overview", heading_style))
        overview_data = [
            ["Total Candidates Analyzed", str(len(coverage))],
            ["Evaluation Criteria", str(len([c for c in coverage.columns if c not in ('Candidate', 'Overall')]))],
            ["Highest Score Achieved", f"{coverage['Overall'].max()*100:.0f}%"],
            ["Average Score", f"{coverage['Overall'].mean()*100:.0f}%"]
        ]
        overview_table = Table(overview_data, colWidths=[10*cm, 6*cm])
        overview_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#E7F3FF')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(overview_table)
        story.append(Spacer(1, 0.5*cm))
        
        # Top 5 Candidates with color-coded pill ratings
        story.append(Paragraph("Top 5 Candidates", heading_style))
        top5 = coverage[['Candidate', 'Overall']].head(5)
        
        # Create custom paragraph style for pills
        from reportlab.platypus import Paragraph as RLParagraph
        
        top_data = [["Rank", "Candidate Name", "Overall Score", "Rating"]]
        for idx, (_, row) in enumerate(top5.iterrows(), 1):
            score = row['Overall']
            score_pct = f"{score*100:.0f}%"
            
            # Create color-coded pill HTML
            if score >= hi:
                rating_pill = '<para alignment="center" backColor="#10B981" textColor="white" borderPadding="4" borderRadius="12" fontSize="10"><b>STRONG</b></para>'
                bg_color = colors.HexColor('#D1FAE5')
            elif score >= lo:
                rating_pill = '<para alignment="center" backColor="#F59E0B" textColor="white" borderPadding="4" borderRadius="12" fontSize="10"><b>MODERATE</b></para>'
                bg_color = colors.HexColor('#FEF3C7')
            else:
                rating_pill = '<para alignment="center" backColor="#EF4444" textColor="white" borderPadding="4" borderRadius="12" fontSize="10"><b>WEAK</b></para>'
                bg_color = colors.HexColor('#FEE2E2')
            
            top_data.append([str(idx), row['Candidate'], score_pct, Paragraph(rating_pill, styles['Normal'])])
        
        top_table = Table(top_data, colWidths=[2*cm, 7*cm, 3.5*cm, 3.5*cm])
        top_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F77B4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F9FAFB')])
        ]))
        
        # Add background color to rating cells
        for i, (_, row) in enumerate(top5.iterrows(), 1):
            score = row['Overall']
            if score >= hi:
                bg_color = colors.HexColor('#D1FAE5')
            elif score >= lo:
                bg_color = colors.HexColor('#FEF3C7')
            else:
                bg_color = colors.HexColor('#FEE2E2')
            top_table.setStyle(TableStyle([('BACKGROUND', (3, i), (3, i), bg_color)]))
        
        story.append(top_table)
        story.append(Spacer(1, 0.5*cm))
        
        # AI Insights for Top 5 (if available) - with horizontal separators
        if insights:
            story.append(Paragraph("Key Insights for Top Candidates", heading_style))
            
            # Create bullet style with proper indentation
            bullet_style = ParagraphStyle(
                'BulletStyle',
                parent=styles['Normal'],
                leftIndent=20,
                firstLineIndent=-20,
                spaceBefore=3,
                spaceAfter=3,
                bulletIndent=0
            )
            
            for idx, candidate_name in enumerate(top5['Candidate'].head(5), 1):
                cand_insights = insights.get(candidate_name, {})
                
                if cand_insights:
                    story.append(Paragraph(f"<b>{idx}. {candidate_name}</b>", styles['Heading3']))
                    
                    # Strengths
                    strengths = cand_insights.get('top', [])
                    if strengths:
                        story.append(Paragraph("<b>Strengths:</b>", styles['Normal']))
                        for s in strengths[:3]:  # Top 3 strengths only
                            story.append(Paragraph(f"• {s}", bullet_style))
                        story.append(Spacer(1, 0.2*cm))  # Space after Strengths
                    
                    # Gaps
                    gaps = cand_insights.get('gaps', [])
                    if gaps:
                        story.append(Paragraph("<b>Development Areas:</b>", styles['Normal']))
                        for g in gaps[:3]:  # Top 3 gaps only
                            story.append(Paragraph(f"• {g}", bullet_style))
                        story.append(Spacer(1, 0.2*cm))  # Space after Development Areas
                    
                    # Notes with heading
                    notes = cand_insights.get('notes', '')
                    if notes:
                        story.append(Paragraph("<b>Overall Assessment:</b>", styles['Normal']))
                        story.append(Paragraph(f"{notes}", styles['Normal']))
                    
                    # Add horizontal separator between candidates (except last one)
                    if idx < len(top5['Candidate'].head(5)):
                        story.append(Spacer(1, 0.2*cm))
                        from reportlab.platypus import HRFlowable
                        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E5E7EB')))
                        story.append(Spacer(1, 0.3*cm))
                    else:
                        story.append(Spacer(1, 0.3*cm))
                else:
                    story.append(Paragraph(f"<b>{idx}. {candidate_name}</b>", styles['Heading3']))
                    story.append(Paragraph("<i>No AI insights generated for this candidate</i>", styles['Normal']))
                    
                    # Add horizontal separator between candidates (except last one)
                    if idx < len(top5['Candidate'].head(5)):
                        story.append(Spacer(1, 0.2*cm))
                        from reportlab.platypus import HRFlowable
                        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E5E7EB')))
                        story.append(Spacer(1, 0.3*cm))
                    else:
                        story.append(Spacer(1, 0.3*cm))
        
        # Recommendation (intelligent, context-aware - matching Streamlit)
        story.append(PageBreak())
        story.append(Paragraph("Recommendation", heading_style))
        
        top_candidate_name = coverage.iloc[0]['Candidate']
        top_score = coverage.iloc[0]['Overall']
        
        # Check if we have multiple strong candidates
        strong_candidates = coverage[coverage['Overall'] >= hi]
        moderate_candidates = coverage[(coverage['Overall'] >= lo) & (coverage['Overall'] < hi)]
        weak_candidates = coverage[coverage['Overall'] < lo]
        
        # Build intelligent recommendation based on actual data
        recommendation_parts = []
        
        if len(strong_candidates) == 0:
            # No strong candidates at all
            if len(moderate_candidates) > 0:
                recommendation_parts.append(
                    f"<b>Caution:</b> No candidates achieved a strong match score (≥{hi*100:.0f}%). "
                    f"The highest score was <b>{top_score*100:.0f}%</b> for <b>{top_candidate_name}</b>, "
                    f"indicating a <b>moderate match</b>."
                )
                recommendation_parts.append(
                    f"We recommend carefully reviewing the {len(moderate_candidates)} moderate-scoring candidate(s) "
                    f"to identify specific skill gaps, or consider expanding the candidate pool."
                )
            else:
                recommendation_parts.append(
                    f"<b>Warning:</b> All candidates scored below the moderate threshold ({lo*100:.0f}%). "
                    f"The highest score was only <b>{top_score*100:.0f}%</b> for <b>{top_candidate_name}</b>."
                )
                recommendation_parts.append(
                    "We recommend reconsidering the job requirements or sourcing additional candidates, "
                    "as the current pool shows weak alignment with the position criteria."
                )
        
        elif len(strong_candidates) == 1:
            # Clear winner
            recommendation_parts.append(
                f"<b>{top_candidate_name}</b> is the clear leading candidate with a strong overall score "
                f"of <b>{top_score*100:.0f}%</b>, demonstrating excellent alignment with the position requirements."
            )
            
            if len(moderate_candidates) > 0:
                recommendation_parts.append(
                    f"Additionally, {len(moderate_candidates)} candidate(s) achieved moderate scores and could serve as backup options."
                )
            
            recommendation_parts.append(
                f"We recommend prioritizing <b>{top_candidate_name}</b> for the next stage of recruitment."
            )
        
        else:
            # Multiple strong candidates
            top_3_strong = strong_candidates.head(3)
            score_range = top_3_strong['Overall'].max() - top_3_strong['Overall'].min()
            
            if score_range < 0.10:  # Very close scores
                names = ", ".join([f"<b>{row['Candidate']}</b>" for _, row in top_3_strong.iterrows()])
                recommendation_parts.append(
                    f"We have <b>{len(strong_candidates)} strong candidates</b> with very similar scores "
                    f"(range: {top_3_strong['Overall'].min()*100:.0f}%–{top_3_strong['Overall'].max()*100:.0f}%). "
                    f"The top candidates are: {names}."
                )
                recommendation_parts.append(
                    "Given the close scoring, we recommend interviewing multiple candidates to assess "
                    "cultural fit, communication skills, and other qualitative factors."
                )
            else:
                recommendation_parts.append(
                    f"<b>{top_candidate_name}</b> is the leading candidate with a score of <b>{top_score*100:.0f}%</b>, "
                    f"followed by {len(strong_candidates)-1} other strong candidate(s)."
                )
                recommendation_parts.append(
                    f"We recommend prioritizing <b>{top_candidate_name}</b>, while keeping other strong candidates "
                    f"as viable alternatives."
                )
        
        # Combine recommendation parts
        for part in recommendation_parts:
            story.append(Paragraph(part, styles['Normal']))
            story.append(Spacer(1, 0.2*cm))
        
        story.append(Spacer(1, 0.3*cm))
        
        # Legend - Fixed typo
        story.append(Paragraph("Score Interpretation", heading_style))
        legend_text = (
            f"• <b>Strong (≥{hi*100:.0f}%):</b> Excellent alignment with requirements<br/>"
            f"• <b>Moderate ({lo*100:.0f}%–{hi*100:.0f}%):</b> Acceptable fit with some gaps<br/>"
            f"• <b>Weak (&lt;{lo*100:.0f}%):</b> Significant gaps in required areas"
        )
        story.append(Paragraph(legend_text, styles['Normal']))
        story.append(Spacer(1, 0.2*cm))
        
        # Build PDF
        doc.build(story)
        buf.seek(0)
        return buf.getvalue()
    
    except Exception as e:
        print(f"PDF generation error: {e}")
        return None


# ============================================================================
# WORD EXPORT - Executive Summary (Formatted Document)
# ============================================================================

def to_executive_summary_word(
    coverage: pd.DataFrame,
    insights: Dict[str, Dict],
    jd_text: str,
    cat_map: Dict[str, str],
    hi: float = 0.75,
    lo: float = 0.35,
    jd_filename: str = "Job Description",
    job_number: Optional[int] = None
) -> Optional[bytes]:
    """
    Generate a professional executive summary Word document.
    
    Includes:
    - Header with Job # and branding
    - Analysis overview & key metrics
    - Top 5 candidates table with color-coded ratings
    - AI insights for candidates
    - Shortlist recommendation (on separate page)
    - Placeholder for company logo
    
    Args:
        coverage: DataFrame with candidate scores
        insights: Dict mapping candidate_name -> {top, gaps, notes}
        jd_text: Job description text
        cat_map: Criterion -> category mapping
        hi: High threshold (default 0.75 = 75%)
        lo: Low threshold (default 0.35 = 35%)
        jd_filename: Name of JD file for display
        job_number: Analysis/Job ID number for header
    
    Returns:
        Bytes of Word document or None if python-docx not available
    """
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        
        # ===== HEADER WITH JOB # AND LOGO PLACEHOLDER =====
        header_section = doc.sections[0]
        header = header_section.header
        
        # Add logo placeholder (right-aligned)
        logo_para = header.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo_run = logo_para.add_run("[COMPANY LOGO HERE]")
        logo_run.font.size = Pt(10)
        logo_run.font.color.rgb = RGBColor(150, 150, 150)
        logo_run.font.italic = True
        
        # Title with Job Number
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("Executive Summary - Candidate Analysis")
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 119, 180)  # Blue
        
        # Job number and title
        if job_number:
            job_title_para = doc.add_paragraph()
            job_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            job_run = job_title_para.add_run(f"Job #{job_number:04d}: {jd_filename}")
            job_run.font.name = 'Arial'
            job_run.font.size = Pt(14)
            job_run.font.bold = True
        else:
            job_title_para = doc.add_paragraph()
            job_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            job_run = job_title_para.add_run(f"Job Position: {jd_filename}")
            job_run.font.name = 'Arial'
            job_run.font.size = Pt(14)
            job_run.font.bold = True
        
        # Generated date (subtle)
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = RGBColor(107, 114, 128)  # Gray
        
        doc.add_paragraph()  # Spacer
        
        # ===== ANALYSIS OVERVIEW =====
        heading = doc.add_paragraph()
        heading_run = heading.add_run("Analysis Overview")
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(31, 119, 180)
        
        # Overview table
        overview_table = doc.add_table(rows=4, cols=2)
        overview_table.style = 'Light Grid Accent 1'
        overview_data = [
            ("Total Candidates Analyzed", str(len(coverage))),
            ("Evaluation Criteria", str(len([c for c in coverage.columns if c not in ('Candidate', 'Overall')]))),
            ("Highest Score Achieved", f"{coverage['Overall'].max()*100:.0f}%"),
            ("Average Score", f"{coverage['Overall'].mean()*100:.0f}%")
        ]
        
        for i, (label, value) in enumerate(overview_data):
            row_cells = overview_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value
            # Bold labels
            for paragraph in row_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        doc.add_paragraph()  # Spacer
        
        # ===== TOP 5 CANDIDATES TABLE WITH COLOR-CODED PILLS =====
        heading = doc.add_paragraph()
        heading_run = heading.add_run("Top 5 Candidates")
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(31, 119, 180)
        
        top5 = coverage[['Candidate', 'Overall']].head(5)
        
        # Create table
        candidates_table = doc.add_table(rows=len(top5)+1, cols=4)
        candidates_table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = candidates_table.rows[0].cells
        headers = ["Rank", "Candidate Name", "Overall Score", "Rating"]
        for i, header_text in enumerate(headers):
            header_cells[i].text = header_text
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # Blue background for header
            from docx.oxml.shared import OxmlElement, qn
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '1F77B4')
            header_cells[i]._element.get_or_add_tcPr().append(shading_elm)
        
        # Data rows with color-coded ratings
        for idx, (_, row) in enumerate(top5.iterrows(), 1):
            score = row['Overall']
            score_pct = f"{score*100:.0f}%"
            
            # Determine rating and color
            if score >= hi:
                rating_text = "STRONG"
                bg_color = 'D1FAE5'  # Light green
            elif score >= lo:
                rating_text = "MODERATE"
                bg_color = 'FEF3C7'  # Light yellow
            else:
                rating_text = "WEAK"
                bg_color = 'FEE2E2'  # Light red
            
            row_cells = candidates_table.rows[idx].cells
            row_cells[0].text = str(idx)
            row_cells[1].text = row['Candidate']
            row_cells[2].text = score_pct
            row_cells[3].text = rating_text
            
            # Center align all cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Color the rating cell
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), bg_color)
            row_cells[3]._element.get_or_add_tcPr().append(shading_elm)
            
            # Bold the rating text
            for paragraph in row_cells[3].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        doc.add_paragraph()  # Spacer
        
        # ===== KEY INSIGHTS FOR TOP CANDIDATES =====
        if insights:
            heading = doc.add_paragraph()
            heading_run = heading.add_run("Key Insights for Top Candidates")
            heading_run.font.name = 'Arial'
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(31, 119, 180)
            
            for idx, candidate_name in enumerate(top5['Candidate'].head(5), 1):
                cand_insights = insights.get(candidate_name, {})
                
                # Candidate name
                cand_heading = doc.add_paragraph()
                cand_run = cand_heading.add_run(f"{idx}. {candidate_name}")
                cand_run.font.name = 'Arial'
                cand_run.font.size = Pt(12)
                cand_run.font.bold = True
                
                if cand_insights:
                    # Strengths
                    strengths = cand_insights.get('top', [])
                    if strengths:
                        strengths_para = doc.add_paragraph()
                        strengths_run = strengths_para.add_run("Strengths:")
                        strengths_run.font.bold = True
                        for s in strengths[:3]:
                            doc.add_paragraph(s, style='List Bullet')
                    
                    # Gaps/Development Areas
                    gaps = cand_insights.get('gaps', [])
                    if gaps:
                        gaps_para = doc.add_paragraph()
                        gaps_run = gaps_para.add_run("Development Areas:")
                        gaps_run.font.bold = True
                        for g in gaps[:3]:
                            doc.add_paragraph(g, style='List Bullet')
                    
                    # Overall Assessment with heading
                    notes = cand_insights.get('notes', '')
                    if notes:
                        assessment_para = doc.add_paragraph()
                        assessment_run = assessment_para.add_run("Overall Assessment:")
                        assessment_run.font.bold = True
                        notes_para = doc.add_paragraph()
                        notes_run = notes_para.add_run(notes)
                else:
                    no_insights = doc.add_paragraph("No AI insights generated for this candidate")
                    for run in no_insights.runs:
                        run.font.italic = True
                
                # Add horizontal separator (except for last candidate)
                if idx < len(top5['Candidate'].head(5)):
                    doc.add_paragraph('_' * 80)  # Simple text separator
        
        # ===== PAGE BREAK BEFORE RECOMMENDATION =====
        doc.add_page_break()
        
        # ===== RECOMMENDATION =====
        heading = doc.add_paragraph()
        heading_run = heading.add_run("Recommendation")
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(31, 119, 180)
        
        top_candidate_name = coverage.iloc[0]['Candidate']
        top_score = coverage.iloc[0]['Overall']
        
        # Build intelligent recommendation (same logic as PDF)
        strong_candidates = coverage[coverage['Overall'] >= hi]
        moderate_candidates = coverage[(coverage['Overall'] >= lo) & (coverage['Overall'] < hi)]
        
        recommendation_parts = []
        
        if len(strong_candidates) == 0:
            if len(moderate_candidates) > 0:
                recommendation_parts.append(
                    f"Caution: No candidates achieved a strong match score (≥{hi*100:.0f}%). "
                    f"The highest score was {top_score*100:.0f}% for {top_candidate_name}, indicating a moderate match."
                )
                recommendation_parts.append(
                    f"We recommend carefully reviewing the {len(moderate_candidates)} moderate-scoring candidate(s) "
                    f"to identify specific skill gaps, or consider expanding the candidate pool."
                )
            else:
                recommendation_parts.append(
                    f"Warning: All candidates scored below the moderate threshold ({lo*100:.0f}%). "
                    f"The highest score was only {top_score*100:.0f}% for {top_candidate_name}."
                )
                recommendation_parts.append(
                    "We recommend reconsidering the job requirements or sourcing additional candidates, "
                    "as the current pool shows weak alignment with the position criteria."
                )
        elif len(strong_candidates) == 1:
            recommendation_parts.append(
                f"{top_candidate_name} is the clear leading candidate with a strong overall score "
                f"of {top_score*100:.0f}%, demonstrating excellent alignment with the position requirements."
            )
            if len(moderate_candidates) > 0:
                recommendation_parts.append(
                    f"Additionally, {len(moderate_candidates)} candidate(s) achieved moderate scores and could serve as backup options."
                )
            recommendation_parts.append(
                f"We recommend prioritizing {top_candidate_name} for the next stage of recruitment."
            )
        else:
            top_3_strong = strong_candidates.head(3)
            score_range = top_3_strong['Overall'].max() - top_3_strong['Overall'].min()
            
            if score_range < 0.10:
                names = ", ".join([row['Candidate'] for _, row in top_3_strong.iterrows()])
                recommendation_parts.append(
                    f"We have {len(strong_candidates)} strong candidates with very similar scores "
                    f"(range: {top_3_strong['Overall'].min()*100:.0f}%–{top_3_strong['Overall'].max()*100:.0f}%). "
                    f"The top candidates are: {names}."
                )
                recommendation_parts.append(
                    "Given the close scoring, we recommend interviewing multiple candidates to assess "
                    "cultural fit, communication skills, and other qualitative factors."
                )
            else:
                recommendation_parts.append(
                    f"{top_candidate_name} is the leading candidate with a score of {top_score*100:.0f}%, "
                    f"followed by {len(strong_candidates)-1} other strong candidate(s)."
                )
                recommendation_parts.append(
                    f"We recommend prioritizing {top_candidate_name}, while keeping other strong candidates as viable alternatives."
                )
        
        for part in recommendation_parts:
            doc.add_paragraph(part)
        
        doc.add_paragraph()  # Spacer
        
        # ===== SCORE INTERPRETATION LEGEND =====
        heading = doc.add_paragraph()
        heading_run = heading.add_run("Score Interpretation")
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(31, 119, 180)
        
        doc.add_paragraph(f"• Strong (≥{hi*100:.0f}%): Excellent alignment with requirements")
        doc.add_paragraph(f"• Moderate ({lo*100:.0f}%–{hi*100:.0f}%): Acceptable fit with some gaps")
        doc.add_paragraph(f"• Weak (<{lo*100:.0f}%): Significant gaps in required areas")
        
        # Save to bytes
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    
    except Exception as e:
        print(f"Word generation error: {e}")
        return None


# End of export_utils.py
