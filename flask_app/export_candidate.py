"""
Individual candidate export functions - Part 2 of export utilities
Contains PDF and Word export functions for individual candidate reports
"""

from typing import Dict, List, Any, Optional, Tuple
import io
import pandas as pd
from datetime import datetime

# Import availability flags from main export_utils
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ============================================================================
# PDF EXPORT - Individual Candidate Report
# ============================================================================

def to_individual_candidate_pdf(
    candidate_name: str,
    coverage_row: pd.Series,
    insights: Dict[str, Any],
    evidence_map: Dict[Tuple[str,str], Tuple[str,float]],
    cat_map: Dict[str, str],
    hi: float = 0.75,
    lo: float = 0.35,
    include_evidence: bool = False,
    job_title: str = "",
    gpt_candidates: List[str] = None,
    job_number: int = None
) -> Optional[bytes]:
    """
    Generate individual candidate report PDF with disclaimer for non-AI candidates.
    
    Args:
        candidate_name: Full name of candidate
        coverage_row: Pandas Series with scores for all criteria + Overall
        insights: Dict with 'top', 'gaps', 'notes' keys
        evidence_map: Dict[(candidate, criterion)] -> (snippet, score)
        cat_map: Criterion -> category mapping
        hi: High threshold
        lo: Low threshold
        include_evidence: Whether to include evidence snippets
        job_title: Job title for header
        gpt_candidates: List of candidates that had AI insights generated
        job_number: Job number for header (e.g., 18 for "Job #0018")
    
    Returns:
        Bytes of PDF or None if reportlab not available
    """
    if not REPORTLAB_AVAILABLE:
        return None
    
    try:
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#1F77B4'),
            spaceAfter=20,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#1F77B4'),
            spaceAfter=10,
            spaceBefore=10
        )
        
        # Title with Job# in header
        job_header = f"Job #{job_number:04d}: {job_title}" if job_number and job_title else (f"Job #{job_number:04d}" if job_number else (f"Position: {job_title}" if job_title else ""))
        if job_header:
            story.append(Paragraph(job_header, styles['Normal']))
        story.append(Paragraph(f"Candidate Report: {candidate_name}", title_style))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 0.5*cm))
        
        # Overall Score (convert to percentage and bold labels)
        overall_score = coverage_row.get('Overall', 0.0)
        rating = "Strong Match" if overall_score >= hi else ("Moderate Match" if overall_score >= lo else "Weak Match")
        
        score_data = [
            [Paragraph("<b>Overall Score</b>", styles['Normal']), f"{overall_score * 100:.0f}%"],
            [Paragraph("<b>Rating</b>", styles['Normal']), rating]
        ]
        score_table = Table(score_data, colWidths=[8*cm, 8*cm])
        score_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#E7F3FF')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(score_table)
        story.append(Spacer(1, 0.5*cm))
        
        # Check if this candidate has GPT-generated AI insights
        if gpt_candidates is None:
            gpt_candidates = []
        has_ai_insights = candidate_name in gpt_candidates
        
        if not has_ai_insights:
            # Add disclaimer for non-AI analyzed candidates
            disclaimer_style = ParagraphStyle(
                'Disclaimer',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#856404'),
                backColor=colors.HexColor('#FFF3CD'),
                borderColor=colors.HexColor('#856404'),
                borderWidth=1,
                borderPadding=10,
                spaceAfter=15,
                spaceBefore=5
            )
            story.append(Paragraph(
                "<b>ℹ️ Note:</b> This candidate was not included in the GPT Insights analysis. "
                "The strengths and development areas shown below are based on scoring analysis only "
                "(top 3 and bottom 3 criteria by match score). "
                "To receive AI-generated narrative insights, include this candidate in your 'GPT Insights' "
                "selection and re-run the analysis.",
                disclaimer_style
            ))
            story.append(Spacer(1, 0.3*cm))
        
        # Key Strengths
        strengths = insights.get('top', []) if insights else []
        story.append(Paragraph("Key Strengths", heading_style))
        if strengths:
            for s in strengths:
                story.append(Paragraph(f"• {s}", styles['Normal']))
        else:
            story.append(Paragraph("<i>No data available</i>", styles['Normal']))
        story.append(Spacer(1, 0.3*cm))
        
        # Development Areas / Gaps
        gaps = insights.get('gaps', []) if insights else []
        story.append(Paragraph("Development Areas", heading_style))
        if gaps:
            for g in gaps:
                story.append(Paragraph(f"• {g}", styles['Normal']))
        else:
            story.append(Paragraph("<i>No data available</i>", styles['Normal']))
        story.append(Spacer(1, 0.3*cm))
        
        # Additional Notes
        notes = insights.get('notes', '')
        if notes:
            story.append(Paragraph("Additional Notes", heading_style))
            story.append(Paragraph(notes, styles['Normal']))
            story.append(Spacer(1, 0.3*cm))
        
        # Detailed Scores by Criteria
        story.append(Paragraph("Detailed Scores by Criteria", heading_style))
        
        # Group by category
        criteria_by_cat = {}
        for col in coverage_row.index:
            if col not in ('Candidate', 'Overall'):
                cat = cat_map.get(col, 'Uncategorized')
                if cat not in criteria_by_cat:
                    criteria_by_cat[cat] = []
                criteria_by_cat[cat].append((col, coverage_row[col]))
        
        for category in sorted(criteria_by_cat.keys()):
            story.append(Paragraph(f"<b>{category}</b>", styles['Heading4']))
            
            criteria_data = [["Criterion", "Score", "Status"]]
            for crit, score in sorted(criteria_by_cat[category], key=lambda x: x[1], reverse=True):
                # Use Paragraph for criterion name to enable text wrapping
                crit_para = Paragraph(crit, styles['Normal'])
                # Color-coded pill instead of text rating
                if score >= hi:
                    status_pill = Paragraph('<para align="center" backColor="#D1FAE5" borderColor="#10B981" borderWidth="1" borderPadding="3"><font color="#065F46" size="9"><b>STRONG</b></font></para>', styles['Normal'])
                elif score >= lo:
                    status_pill = Paragraph('<para align="center" backColor="#FEF3C7" borderColor="#F59E0B" borderWidth="1" borderPadding="3"><font color="#92400E" size="9"><b>MODERATE</b></font></para>', styles['Normal'])
                else:
                    status_pill = Paragraph('<para align="center" backColor="#FEE2E2" borderColor="#EF4444" borderWidth="1" borderPadding="3"><font color="#991B1B" size="9"><b>WEAK</b></font></para>', styles['Normal'])
                criteria_data.append([crit_para, f"{score * 100:.0f}%", status_pill])
            
            crit_table = Table(criteria_data, colWidths=[10*cm, 3*cm, 3*cm])
            crit_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E7F3FF')),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
            ]))
            

            story.append(crit_table)
            story.append(Spacer(1, 0.3*cm))
        
        # Evidence snippets (optional)
        if include_evidence and evidence_map:
            story.append(Paragraph("Evidence Snippets (Selected)", heading_style))
            story.append(Paragraph("<i>Showing evidence for top-scored criteria</i>", styles['Normal']))
            
            # Get top 5 criteria by score
            top_criteria = sorted(
                [(col, coverage_row[col]) for col in coverage_row.index if col not in ('Candidate', 'Overall')],
                key=lambda x: x[1],
                reverse=True
            )[:5]
            
            for crit, score in top_criteria:
                evidence_key = (candidate_name, crit)
                if evidence_key in evidence_map:
                    evidence_tuple = evidence_map[evidence_key]
                    # Handle both old format (snippet, score) and new format (snippet, score, density)
                    snippet = evidence_tuple[0]
                    # Clean and truncate snippet, ensuring we get unique content
                    clean_snippet = snippet.strip()
                    if len(clean_snippet) > 300:
                        # Try to find a sentence boundary near 300 chars
                        truncate_pos = clean_snippet.rfind('.', 200, 300)
                        if truncate_pos == -1:
                            truncate_pos = 300
                        clean_snippet = clean_snippet[:truncate_pos+1]
                    story.append(Paragraph(f"<b>{crit}</b> (Score: {score * 100:.0f}%)", styles['Heading4']))
                    story.append(Paragraph(f"<i>{clean_snippet}</i>", styles['Normal']))
                    story.append(Spacer(1, 0.2*cm))
        
        # Build PDF
        doc.build(story)
        buf.seek(0)
        pdf_bytes = buf.getvalue()
        print(f"✓ Individual PDF built successfully, size: {len(pdf_bytes)} bytes")
        return pdf_bytes
    
    except Exception as e:
        print(f"✗ PDF generation error for {candidate_name}: {e}")
        import traceback
        traceback.print_exc()
        return None


# ============================================================================
# WORD EXPORT - Individual Candidate Report
# ============================================================================

def to_individual_candidate_docx(
    candidate_name: str,
    coverage_row: pd.Series,
    insights: Dict[str, Any],
    evidence_map: Dict[Tuple[str,str], Tuple[str,float]],
    cat_map: Dict[str, str],
    hi: float = 0.75,
    lo: float = 0.35,
    include_evidence: bool = False,
    job_title: str = "",
    gpt_candidates: List[str] = None,
    job_number: int = None
) -> Optional[bytes]:
    """
    Generate editable Word document for individual candidate.
    
    Same structure as PDF but fully editable for customization.
    Includes disclaimer for candidates without AI insights.
    
    Args:
        candidate_name: Full name of candidate
        coverage_row: Pandas Series with scores
        insights: Dict with 'top', 'gaps', 'notes'
        cat_map: Criterion -> category mapping
        hi: High threshold
        lo: Low threshold
        job_title: Job title for header
        gpt_candidates: List of AI-analyzed candidates
    
    Returns:
        Bytes of Word document or None if python-docx not available
    """
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        
        # Title with Job# in header
        if job_number and job_title:
            job_para = doc.add_paragraph(f"Job #{job_number:04d}: {job_title}")
            job_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif job_number:
            job_para = doc.add_paragraph(f"Job #{job_number:04d}")
            job_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif job_title:
            position_para = doc.add_paragraph(f"Position: {job_title}")
            position_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title = doc.add_heading(f'Candidate Report: {candidate_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Overall Score (convert to percentage and bold labels)
        overall_score = coverage_row.get('Overall', 0.0)
        rating = "Strong Match" if overall_score >= hi else ("Moderate Match" if overall_score >= lo else "Weak Match")
        
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Light Grid Accent 1'
        cell = table.cell(0, 0)
        cell.text = 'Overall Score'
        cell.paragraphs[0].runs[0].bold = True
        table.cell(0, 1).text = f"{overall_score * 100:.0f}%"
        cell = table.cell(1, 0)
        cell.text = 'Rating'
        cell.paragraphs[0].runs[0].bold = True
        table.cell(1, 1).text = rating
        
        doc.add_paragraph()  # Spacing
        
        # Check if this candidate has GPT-generated AI insights
        if gpt_candidates is None:
            gpt_candidates = []
        has_ai_insights = candidate_name in gpt_candidates
        
        if not has_ai_insights:
            # Add disclaimer for non-AI analyzed candidates
            disclaimer_para = doc.add_paragraph()
            disclaimer_para.add_run('ℹ️ Note: ').bold = True
            disclaimer_para.add_run(
                'This candidate was not included in the GPT Insights analysis. '
                'The strengths and development areas shown below are based on scoring analysis only '
                '(top 3 and bottom 3 criteria by match score). '
                'To receive AI-generated narrative insights, include this candidate in your '
                '\'GPT Insights\' selection and re-run the analysis.'
            )
            # Style the disclaimer
            for run in disclaimer_para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(133, 100, 4)
            doc.add_paragraph()  # Spacing
        
        # Key Strengths
        strengths = insights.get('top', []) if insights else []
        doc.add_heading('Key Strengths', 1)
        if strengths:
            for s in strengths:
                doc.add_paragraph(s, style='List Bullet')
        else:
            doc.add_paragraph('No data available', style='Intense Quote')
        
        # Development Areas
        gaps = insights.get('gaps', []) if insights else []
        doc.add_heading('Development Areas', 1)
        if gaps:
            for g in gaps:
                doc.add_paragraph(g, style='List Bullet')
        else:
            doc.add_paragraph('No data available', style='Intense Quote')
        
        # Additional Notes
        notes = insights.get('notes', '')
        if notes:
            doc.add_heading('Additional Notes', 1)
            doc.add_paragraph(notes)
        
        # Detailed Scores by Criteria
        doc.add_heading('Detailed Scores by Criteria', 1)
        
        # Group by category
        criteria_by_cat = {}
        for col in coverage_row.index:
            if col not in ('Candidate', 'Overall'):
                cat = cat_map.get(col, 'Uncategorized')
                if cat not in criteria_by_cat:
                    criteria_by_cat[cat] = []
                criteria_by_cat[cat].append((col, coverage_row[col]))
        
        for category in sorted(criteria_by_cat.keys()):
            doc.add_heading(category, 2)
            
            sorted_criteria = sorted(criteria_by_cat[category], key=lambda x: x[1], reverse=True)
            
            table = doc.add_table(rows=len(sorted_criteria)+1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Header
            table.cell(0, 0).text = 'Criterion'
            table.cell(0, 1).text = 'Score'
            table.cell(0, 2).text = 'Status'
            
            # Data with percentage scores and color-coded pills
            for idx, (crit, score) in enumerate(sorted_criteria, 1):
                table.cell(idx, 0).text = crit
                table.cell(idx, 1).text = f"{score * 100:.0f}%"
                
                # Add color-coded status pill
                status_cell = table.cell(idx, 2)
                if score >= hi:
                    status_cell.text = 'STRONG'
                    # Set green background
                    from docx.oxml import OxmlElement
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'D1FAE5')
                    status_cell._element.get_or_add_tcPr().append(shading_elm)
                    status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(6, 95, 70)
                    status_cell.paragraphs[0].runs[0].bold = True
                elif score >= lo:
                    status_cell.text = 'MODERATE'
                    # Set amber background
                    from docx.oxml import OxmlElement
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FEF3C7')
                    status_cell._element.get_or_add_tcPr().append(shading_elm)
                    status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(146, 64, 14)
                    status_cell.paragraphs[0].runs[0].bold = True
                else:
                    status_cell.text = 'WEAK'
                    # Set red background
                    from docx.oxml import OxmlElement
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FEE2E2')
                    status_cell._element.get_or_add_tcPr().append(shading_elm)
                    status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(153, 27, 27)
                    status_cell.paragraphs[0].runs[0].bold = True
                # Center align status
                status_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Evidence snippets (optional) - matching PDF structure
        if include_evidence and evidence_map:
            doc.add_heading('Detailed Scores & Evidence', 1)
            doc.add_paragraph('Showing evidence for top-scored criteria', style='Intense Quote')
            
            # Get top 5 criteria by score
            all_criteria = [(col, coverage_row[col]) for col in coverage_row.index if col not in ('Candidate', 'Overall')]
            top_criteria = sorted(all_criteria, key=lambda x: x[1], reverse=True)[:5]
            
            for crit, score in top_criteria:
                evidence_key = (candidate_name, crit)
                if evidence_key in evidence_map:
                    evidence_tuple = evidence_map[evidence_key]
                    # Handle both old format (snippet, score) and new format (snippet, score, density)
                    snippet = evidence_tuple[0]
                    # Clean and truncate snippet
                    clean_snippet = snippet.strip()
                    if len(clean_snippet) > 300:
                        truncate_pos = clean_snippet.rfind('.', 200, 300)
                        if truncate_pos == -1:
                            truncate_pos = 300
                        clean_snippet = clean_snippet[:truncate_pos+1]
                    
                    doc.add_heading(f"{crit} (Score: {score * 100:.0f}%)", 2)
                    para = doc.add_paragraph(clean_snippet)
                    para.style = 'Intense Quote'
        
        # Save to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    
    except Exception as e:
        print(f"Word document generation error: {e}")
        return None


# ============================================================================
# WORD EXPORT - Executive Summary
# ============================================================================

def to_executive_summary_docx(
    coverage: pd.DataFrame,
    insights: Dict[str, Dict],
    jd_text: str,
    cat_map: Dict[str, str],
    hi: float = 0.75,
    lo: float = 0.35,
    jd_filename: str = "Job Description"
) -> Optional[bytes]:
    """
    Generate editable Word document for executive summary.
    
    Mirrors the PDF executive summary but fully editable.
    
    Args:
        coverage: DataFrame with candidate scores
        insights: Dict mapping candidate -> insights
        jd_text: Job description text
        cat_map: Criterion -> category mapping
        hi: High threshold
        lo: Low threshold
        jd_filename: JD file name for display
    
    Returns:
        Bytes of Word document or None if python-docx not available
    """
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Executive Summary - Candidate Analysis', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        job_para = doc.add_paragraph(f"Job Position: {jd_filename}")
        job_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Analysis Overview
        doc.add_heading('Analysis Overview', 1)
        
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.cell(0, 0).text = 'Total Candidates Analyzed'
        table.cell(0, 1).text = str(len(coverage))
        
        table.cell(1, 0).text = 'Evaluation Criteria'
        table.cell(1, 1).text = str(len([c for c in coverage.columns if c not in ('Candidate', 'Overall')]))
        
        table.cell(2, 0).text = 'Highest Score Achieved'
        table.cell(2, 1).text = f"{coverage['Overall'].max():.2f}"
        
        table.cell(3, 0).text = 'Average Score'
        table.cell(3, 1).text = f"{coverage['Overall'].mean():.2f}"
        
        doc.add_paragraph()
        
        # Top 5 Candidates
        doc.add_heading('Top 5 Candidates', 1)
        
        top5 = coverage[['Candidate', 'Overall']].head(5)
        
        table = doc.add_table(rows=len(top5)+1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Header
        table.cell(0, 0).text = 'Rank'
        table.cell(0, 1).text = 'Candidate Name'
        table.cell(0, 2).text = 'Overall Score'
        table.cell(0, 3).text = 'Rating'
        
        # Data
        for idx, (_, row) in enumerate(top5.iterrows(), 1):
            score = row['Overall']
            rating = "Strong Match" if score >= hi else ("Moderate Match" if score >= lo else "Weak Match")
            table.cell(idx, 0).text = str(idx)
            table.cell(idx, 1).text = row['Candidate']
            table.cell(idx, 2).text = f"{score:.2f}"
            table.cell(idx, 3).text = rating
        
        doc.add_paragraph()
        
        # AI Insights for Top 5
        if insights:
            doc.add_heading('Key Insights for Top Candidates', 1)
            
            for idx, candidate_name in enumerate(top5['Candidate'].head(5), 1):
                cand_insights = insights.get(candidate_name, {})
                
                doc.add_heading(f"{idx}. {candidate_name}", 2)
                
                if cand_insights:
                    # Strengths
                    strengths = cand_insights.get('top', [])
                    if strengths:
                        doc.add_paragraph('Strengths:', style='Heading 4')
                        for s in strengths[:3]:
                            doc.add_paragraph(s, style='List Bullet')
                    
                    # Gaps
                    gaps = cand_insights.get('gaps', [])
                    if gaps:
                        doc.add_paragraph('Development Areas:', style='Heading 4')
                        for g in gaps[:3]:
                            doc.add_paragraph(g, style='List Bullet')
                    
                    # Notes
                    notes = cand_insights.get('notes', '')
                    if notes:
                        doc.add_paragraph(notes, style='Intense Quote')
                else:
                    doc.add_paragraph('No AI insights generated for this candidate', style='Intense Quote')
        
        # Recommendation (intelligent, context-aware - matching Streamlit)
        doc.add_page_break()
        doc.add_heading('Recommendation', 1)
        
        top_candidate_name = coverage.iloc[0]['Candidate']
        top_score = coverage.iloc[0]['Overall']
        
        # Check if we have multiple strong candidates
        strong_candidates = coverage[coverage['Overall'] >= hi]
        moderate_candidates = coverage[(coverage['Overall'] >= lo) & (coverage['Overall'] < hi)]
        weak_candidates = coverage[coverage['Overall'] < lo]
        
        # Build intelligent recommendation based on actual data
        if len(strong_candidates) == 0:
            # No strong candidates at all
            if len(moderate_candidates) > 0:
                para1 = doc.add_paragraph()
                para1.add_run('Caution: ').bold = True
                para1.add_run(f'No candidates achieved a strong match score (≥{hi:.2f}). ')
                para1.add_run(f'The highest score was ')
                para1.add_run(f'{top_score:.2f}').bold = True
                para1.add_run(f' for ')
                para1.add_run(top_candidate_name).bold = True
                para1.add_run(', indicating a ')
                para1.add_run('moderate match').bold = True
                para1.add_run('.')
                
                para2 = doc.add_paragraph()
                para2.add_run(f'We recommend carefully reviewing the {len(moderate_candidates)} moderate-scoring candidate(s) ')
                para2.add_run('to identify specific skill gaps, or consider expanding the candidate pool.')
            else:
                para1 = doc.add_paragraph()
                para1.add_run('Warning: ').bold = True
                para1.add_run(f'All candidates scored below the moderate threshold ({lo:.2f}). ')
                para1.add_run(f'The highest score was only ')
                para1.add_run(f'{top_score:.2f}').bold = True
                para1.add_run(f' for ')
                para1.add_run(top_candidate_name).bold = True
                para1.add_run('.')
                
                para2 = doc.add_paragraph()
                para2.add_run('We recommend reconsidering the job requirements or sourcing additional candidates, ')
                para2.add_run('as the current pool shows weak alignment with the position criteria.')
        
        elif len(strong_candidates) == 1:
            # Clear winner
            para1 = doc.add_paragraph()
            para1.add_run(top_candidate_name).bold = True
            para1.add_run(' is the clear leading candidate with a strong overall score of ')
            para1.add_run(f'{top_score:.2f}').bold = True
            para1.add_run(', demonstrating excellent alignment with the position requirements.')
            
            if len(moderate_candidates) > 0:
                para2 = doc.add_paragraph()
                para2.add_run(f'Additionally, {len(moderate_candidates)} candidate(s) achieved moderate scores and could serve as backup options.')
            
            para3 = doc.add_paragraph()
            para3.add_run('We recommend prioritizing ')
            para3.add_run(top_candidate_name).bold = True
            para3.add_run(' for the next stage of recruitment.')
        
        else:
            # Multiple strong candidates
            top_3_strong = strong_candidates.head(3)
            score_range = top_3_strong['Overall'].max() - top_3_strong['Overall'].min()
            
            if score_range < 0.10:  # Very close scores
                names = ", ".join([row['Candidate'] for _, row in top_3_strong.iterrows()])
                para1 = doc.add_paragraph()
                para1.add_run('We have ')
                para1.add_run(f'{len(strong_candidates)} strong candidates').bold = True
                para1.add_run(' with very similar scores ')
                para1.add_run(f'(range: {top_3_strong["Overall"].min():.2f}–{top_3_strong["Overall"].max():.2f}). ')
                para1.add_run(f'The top candidates are: {names}.')
                
                para2 = doc.add_paragraph()
                para2.add_run('Given the close scoring, we recommend interviewing multiple candidates to assess ')
                para2.add_run('cultural fit, communication skills, and other qualitative factors.')
            else:
                para1 = doc.add_paragraph()
                para1.add_run(top_candidate_name).bold = True
                para1.add_run(' is the leading candidate with a score of ')
                para1.add_run(f'{top_score:.2f}').bold = True
                para1.add_run(f', followed by {len(strong_candidates)-1} other strong candidate(s).')
                
                para2 = doc.add_paragraph()
                para2.add_run('We recommend prioritizing ')
                para2.add_run(top_candidate_name).bold = True
                para2.add_run(', while keeping other strong candidates as viable alternatives.')
        
        # Legend
        doc.add_heading('Score Interpretation', 2)
        doc.add_paragraph(f'• Strong (≥{hi:.2f}): Excellent alignment with requirements')
        doc.add_paragraph(f'• Moderate ({lo:.2f}–{hi:.2f}): Acceptable fit with some gaps')
        doc.add_paragraph(f'• Weak (<{lo:.2f}): Significant gaps in required areas')
        
        # Save to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    
    except Exception as e:
        print(f"Word document generation error: {e}")
        return None
